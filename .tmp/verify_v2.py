import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx

doc = docx.Document('Todo_Directo_2_Paginas_Sales_Cheat_Sheet_ES_v2.docx')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')

print('\n--- Content Flow ---')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and len(text) > 3:
        is_bold = any(r.bold for r in p.runs if r.text.strip())
        marker = '■' if is_bold else '·'
        print(f'  {marker} {text[:100]}')

    # Check for page break
    for run in p.runs:
        if 'w:br' in run._r.xml and 'type="page"' in run._r.xml:
            print('  ─── PAGE BREAK ───')

print('\n--- Tables ---')
for i, table in enumerate(doc.tables):
    headers = [cell.text for cell in table.rows[0].cells]
    data_rows = len(table.rows) - 1
    print(f'  Table {i+1}: {headers} ({data_rows} data rows)')
