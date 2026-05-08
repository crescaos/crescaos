"""Generate Cresca Sales Cheat Sheet (2 pages) and Intake Form — 3-Tier Model."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = r'c:\Users\12132\Desktop\DigiFacil\DigiFacil'

def make_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2); s.right_margin = Cm(2)
    sn = d.styles['Normal']; sn.font.name = 'Calibri'; sn.font.size = Pt(10)
    sn.paragraph_format.space_after = Pt(2)
    for lvl in [1,2,3]:
        st = d.styles[f'Heading {lvl}']
        st.font.name = 'Calibri'; st.font.bold = True
        st.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
        st.font.size = Pt([0,16,12,10][lvl])
        st.paragraph_format.space_before = Pt([0,8,6,4][lvl])
        st.paragraph_format.space_after = Pt(3)
    return d

def ap(d, text, bold=False, italic=False, sz=None):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic
    if sz: r.font.size = Pt(sz)
    return p

def at(d, headers, rows):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c = t.cell(0,j); c.text = ''; run = c.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(9); run.font.name = 'Calibri'
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = t.cell(i+1,j); c.text = ''; run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9); run.font.name = 'Calibri'
    d.add_paragraph()
    return t

# ==========================================
# SALES CHEAT SHEET
# ==========================================
cs = make_doc()

cs.add_heading('CRESCA — Guía Rápida de Ventas', level=1)
p = cs.add_paragraph()
r = p.add_run('Sistemas de Crecimiento Empresarial')
r.bold = True; r.font.size = Pt(11)

# --- PAGE 1: What you sell ---
cs.add_heading('🎯 QUÉ VENDEMOS (3 Niveles)', level=2)
ap(cs, 'NO somos: agencia de marketing, constructor de webs, ni agencia de chatbots.', bold=True)
ap(cs, 'SÍ somos: Empresa de Sistemas de Crecimiento Empresarial — Infraestructura + Automatización.', bold=True)

cs.add_heading('Paquete 1: STARTER ($9.99–$49.99/mes)', level=3)
ap(cs, 'Para negocios que inician su presencia digital.')
ap(cs, '✅ Incluye: 1 Workflow Lite (auto-respuesta, confirmación, recordatorio, reseña, o follow-up)')
ap(cs, '💰 Setup: $0–$199  |  4 tiers: DIY Lite → Starter Assist → Presence → Pro')

cs.add_heading('Paquete 2: GROWTH ($50–$100/mes)', level=3)
ap(cs, 'Para negocios que necesitan infraestructura + automatización.')
ap(cs, '✅ Incluye: Website/booking + captura de leads + 1–2 Full Workflows')
ap(cs, '💰 Setup: $199–$499  |  3 tiers: Foundation → +1 Workflow → +2 Workflows')

cs.add_heading('Paquete 3: PRO AUTOMATION (Cotización)', level=3)
ap(cs, 'Para negocios con operaciones complejas.')
ap(cs, '✅ Incluye: Auditoría + 4+ workflows + monitoreo + optimización')
ap(cs, '📋 Recomendado: 12 meses')

cs.add_heading('Add-On & Modules', level=3)
ap(cs, '• Marketing Digital: Solo con Growth o Pro (ad spend lo paga el cliente)')
ap(cs, '• Módulo de Reservas: Capacidad universal, incluido en Growth o Pro según necesidad')
ap(cs, '• Workflows adicionales: +1=$15, +2=$25, +5=$60, +8=$90 (solo Growth)')

cs.add_heading('🔍 AUDITORÍA EN 15 MINUTOS', level=2)
at(cs, ['Área', 'Pregunta', 'Diagnóstico → Paquete'], [
    ['Leads', '¿De dónde vienen clientes?', 'Sin presencia → STARTER'],
    ['Velocidad', '¿Cuánto tardan en responder?', 'Lento → Need workflows'],
    ['Ventas', '¿Cómo cierran?', 'Manual/caótico → GROWTH'],
    ['Operaciones', '¿Qué consume más tiempo?', 'Procesos complejos → PRO'],
    ['Retención', '¿Cómo mantienen contacto?', 'Nada → Workflow Lite min'],
])

cs.add_heading('💬 PITCH EN 30 SEGUNDOS', level=2)
ap(cs, '"Soy [nombre] de Cresca. Construimos sistemas digitales para negocios — desde $9.99 al mes. Empezamos con lo básico y escalamos a medida que creces. ¿Tiene 15 minutos para una auditoría gratuita de su negocio?"', italic=True)

# --- PAGE 2: Pricing, Contracts, Commissions, Objections ---
cs.add_heading('📊 OPCIONES DE PAGO', level=2)
at(cs, ['Opción', 'Detalle', 'Incentivo'], [
    ['Mes a mes', 'Sin compromiso', 'Precio estándar'],
    ['6 meses upfront', 'Pago adelantado', '1 mes gratis o descuento setup'],
    ['12 meses upfront', 'Pago adelantado', '2 meses gratis o descuento setup'],
    ['Plan Flex', 'Setup: 50% hoy + 50% en 30 días', 'Facilita el cierre'],
])
ap(cs, 'NO se permite descuento en mensualidad sin aprobación previa.', bold=True)

cs.add_heading('💰 TU COMPENSACIÓN (Rep Fundador ES)', level=2)
at(cs, ['Concepto', 'Detalle'], [
    ['Salario Base', '$600/mes garantizado'],
    ['Setup Commission', '20% de fees cobrados'],
    ['Residual Commission', '6% de MRR cobrado'],
    ['Clawback', '50% setup si cancela en 60 días'],
])
ap(cs, '📈 Residual es acumulativo — crece cada mes con más clientes activos.')

cs.add_heading('🛡️ TOP 6 OBJECIONES', level=2)
at(cs, ['Objeción', 'Respuesta Clave'], [
    ['"No tengo dinero"', '"STARTER empieza en $9.99. Plan Flex: 50% setup hoy, 50% en 30 días."'],
    ['"Muy caro"', '"¿Cuánto cuesta NO tener un sistema? Menos que un empleado part-time."'],
    ['"Necesito pensarlo"', '"¿Qué info te falta para decidir hoy?"'],
    ['"Ya tengo web"', '"¿Te genera clientes en automático? El valor está en workflows y automatización."'],
    ['"No ahora"', '"¿Cuándo sería buen momento para dejar de perder clientes?"'],
    ['"No confío en IA"', '"La automatización no reemplaza, ayuda. Tú siempre tienes el control."'],
])

cs.add_heading('⚡ FRASES DE CIERRE', level=2)
ap(cs, '• "¿Procedemos con la configuración esta semana?"')
ap(cs, '• "STARTER empieza en $9.99 — menos que un almuerzo."')
ap(cs, '• "Plan Flex: 50% del setup hoy, el resto en 30 días."')
ap(cs, '• "Empiece pequeño y escale cuando crezca."')

ap(cs, '')
ap(cs, 'CRESCA  |  Sistemas de Crecimiento Empresarial  |  Documento Interno — v3.0', bold=True, sz=8)

cs.save(os.path.join(BASE, 'Todo_Directo_2_Paginas_Sales_Cheat_Sheet_ES_v2.docx'))
print("Sales Cheat Sheet saved!")

# ==========================================
# INTAKE FORM
# ==========================================
fi = Document()
for s in fi.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sn = fi.styles['Normal']; sn.font.name = 'Calibri'; sn.font.size = Pt(11)
sn.paragraph_format.space_after = Pt(4)
for lvl in [1,2,3]:
    st = fi.styles[f'Heading {lvl}']
    st.font.name = 'Calibri'; st.font.bold = True
    st.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    st.font.size = Pt([0,16,13,11][lvl])

def afp(text, bold=False, italic=False):
    p = fi.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p

def aft(headers, rows):
    t = fi.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c = t.cell(0,j); c.text = ''; run = c.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.name = 'Calibri'
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = t.cell(i+1,j); c.text = ''; run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10); run.font.name = 'Calibri'
    fi.add_paragraph()
    return t

fi.add_heading('CRESCA — Formulario de Intake', level=1)
afp('Sistemas de Crecimiento Empresarial', bold=True)
afp('Complete este formulario para entender el negocio y preparar la solución ideal.', italic=True)

# Section 1: Contact info
fi.add_heading('1. INFORMACIÓN DEL NEGOCIO', level=2)
aft(['Campo', 'Respuesta'], [
    ['Nombre del negocio', ''],
    ['Nombre del contacto principal', ''],
    ['Cargo / Rol', ''],
    ['Teléfono', ''],
    ['Email', ''],
    ['Dirección del negocio', ''],
    ['Sitio web actual (si tiene)', ''],
    ['Redes sociales (links)', ''],
])

# Section 2: Business type
fi.add_heading('2. TIPO DE NEGOCIO', level=2)
aft(['Pregunta', 'Respuesta'], [
    ['Industria / Sector', ''],
    ['¿Cuántos años tiene el negocio?', ''],
    ['¿Cuántos empleados?', ''],
    ['Rango de ingresos mensuales (aprox)', ''],
    ['¿Quién toma las decisiones?', ''],
])

# Section 3: Service selection — UPDATED
fi.add_heading('3. SELECCIÓN DE SERVICIO', level=2)
afp('Marque el paquete de interés:', bold=True)
aft(['Servicio', 'Seleccione'], [
    ['☐ STARTER — DIY Lite ($9.99/mes)', ''],
    ['☐ STARTER — Starter Assist ($19.99/mes)', ''],
    ['☐ STARTER — Starter Presence ($29.99/mes)', ''],
    ['☐ STARTER — Starter Pro ($49.99/mes)', ''],
    ['☐ GROWTH — Foundation ($50/mes)', ''],
    ['☐ GROWTH — Foundation + 1 Workflow ($75/mes)', ''],
    ['☐ GROWTH — Foundation + 2 Workflows ($100/mes)', ''],
    ['☐ PRO AUTOMATION (Cotización)', ''],
    ['☐ Add-On: Marketing Digital', ''],
])

# Section 3B: Reservas module — NEW
fi.add_heading('3B. MÓDULO DE RESERVAS', level=2)
aft(['Pregunta', 'Respuesta'], [
    ['¿Necesita módulo de reservas?', '☐ Sí  ☐ No'],
    ['Si sí, ¿cuántas propiedades/unidades?', ''],
    ['¿En qué plataformas publica actualmente?', ''],
])

# Section 4: Payment plan — NEW
fi.add_heading('4. PLAN DE PAGO PREFERIDO', level=2)
aft(['Opción', 'Seleccione'], [
    ['☐ Mes a mes (estándar)', ''],
    ['☐ 6 meses por adelantado (1 mes gratis o desc. setup)', ''],
    ['☐ 12 meses por adelantado (2 meses gratis o desc. setup)', ''],
    ['☐ Plan Flex (Setup: 50% hoy + 50% en 30 días)', ''],
])

# Section 5: Audit findings — NEW
fi.add_heading('5. HALLAZGOS DE LA AUDITORÍA (15 Minutos)', level=2)
afp('Completar durante la auditoría con el prospecto:', italic=True)
aft(['Área', 'Hallazgo'], [
    ['Fuentes de leads (¿de dónde vienen clientes?)', ''],
    ['Tiempo de respuesta a prospectos', ''],
    ['Proceso actual de cierre de ventas', ''],
    ['Top 3 dolores/pains operativos', ''],
    ['Resultados de automatización deseados', ''],
    ['¿Tiene sistema de reservas?', '☐ Sí  ☐ No'],
    ['Canales de comunicación con clientes', ''],
])

# Section 6: Digital presence
fi.add_heading('6. PRESENCIA DIGITAL ACTUAL', level=2)
aft(['Pregunta', 'Respuesta'], [
    ['¿Tiene sitio web?', '☐ Sí  ☐ No'],
    ['¿Tiene Google Business Profile?', '☐ Sí  ☐ No'],
    ['¿Usa redes sociales para vender?', '☐ Sí  ☐ No — ¿Cuáles?'],
    ['¿Cómo recibe consultas de clientes?', ''],
    ['¿Usa algún CRM o herramienta de gestión?', ''],
])

# Section 7: Automation discovery
fi.add_heading('7. DESCUBRIMIENTO DE AUTOMATIZACIÓN', level=2)
afp('Solo para clientes de Growth o Pro Automation:', italic=True)
aft(['Pregunta', 'Respuesta'], [
    ['¿Cuáles son sus 3 procesos más repetitivos?', ''],
    ['¿Cuántas horas/semana gasta en tareas administrativas?', ''],
    ['¿Qué software o herramientas usa actualmente?', ''],
    ['¿Tiene procesos de aprobación manual?', ''],
    ['¿Necesita reportes o dashboards?', ''],
    ['¿Qué le gustaría automatizar primero?', ''],
])

# Section 8: Goals
fi.add_heading('8. OBJETIVOS Y EXPECTATIVAS', level=2)
aft(['Pregunta', 'Respuesta'], [
    ['¿Cuál es su objetivo principal al contratar Cresca?', ''],
    ['¿Qué resultado esperaría ver en 3 meses?', ''],
    ['¿Hay alguna fecha límite o evento importante?', ''],
    ['¿Presupuesto mensual disponible para servicios digitales?', ''],
])

# Section 9: Internal use
fi.add_heading('9. USO INTERNO — REPRESENTANTE', level=2)
aft(['Campo', 'Respuesta'], [
    ['Nombre del representante', ''],
    ['Fecha de intake', ''],
    ['Paquete recomendado', '☐ STARTER  ☐ GROWTH  ☐ PRO AUTOMATION'],
    ['Tier/nivel específico', ''],
    ['Módulo de reservas recomendado', '☐ Sí  ☐ No'],
    ['Add-ons recomendados', '☐ Marketing Digital  ☐ Workflows adicionales: ___'],
    ['Plan de pago', '☐ Mensual  ☐ 6 meses  ☐ 12 meses  ☐ Flex'],
    ['Prioridad del lead', '☐ Alta  ☐ Media  ☐ Baja'],
    ['Notas adicionales', ''],
])

afp('')
afp('CRESCA — Sistemas de Crecimiento Empresarial', bold=True)
afp('Formulario de Intake — Documento confidencial para uso interno — v3.0', italic=True)

fi.save(os.path.join(BASE, 'Todo_Directo_Intake_Form.docx'))
print("Intake Form saved!")
print("All done!")
