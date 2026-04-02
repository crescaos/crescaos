import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document()

# ============================================================
# Page setup — Letter size, narrow margins for max content
# ============================================================
section = doc.sections[0]
section.page_width = Cm(21.59)  # Letter
section.page_height = Cm(27.94)
section.top_margin = Cm(1.0)
section.bottom_margin = Cm(1.0)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

# ============================================================
# Style definitions
# ============================================================
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(8.5)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

# Colors
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)
ORANGE = RGBColor(0xE0, 0x6C, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)

def add_shaded_header(doc, text, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="{str(color)}"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def add_section_header(doc, text, color=MEDIUM_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1.2 * indent_level)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(8)
        run = p.add_run(text)
        run.font.size = Pt(8)
    else:
        run = p.add_run(text)
        run.font.size = Pt(8)
    return p

def add_body(doc, text, bold=False, italic=False, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
    
    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(7.5)
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    return table

# ============================================================
# PAGE 1 — PRODUCTS + DECISION TREE + STACK SELLING
# ============================================================

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
run = title.add_run('TODO DIRECTO — GUÍA RÁPIDA DE VENTAS')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = DARK_BLUE

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(6)
run = subtitle.add_run('El Salvador  |  Versión 2.0  |  Febrero 2026  |  Referencia de Campo para Representantes')
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# --- PRODUCT 1: ALTOQUE ---
add_shaded_header(doc, '1️⃣  ALTOQUE APP — Catálogo Digital + Pedidos por WhatsApp')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Ideal para: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Restaurantes, pupuserías, boutiques, panaderías, floristerías, vendedores de Facebook.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Problema: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Pedidos perdidos, desorden en WhatsApp, falta de catálogo profesional.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Posicionamiento: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('"Convierta su WhatsApp en una tienda organizada 24/7."')
run.italic = True
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('IA: ')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = ORANGE
run = p.add_run('Catálogo inteligente, sugerencias de productos, analíticas automáticas.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run('💰 Precios: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Chivo (Gratis) · Pro $14.99/mes · Premium $29.99/mes')
run.font.size = Pt(8)
run.font.color.rgb = MEDIUM_BLUE

# --- PRODUCT 2: WEB ---
add_shaded_header(doc, '2️⃣  SITIO WEB + SISTEMA DE RESERVAS (WaaS)')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Ideal para: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Barberías, clínicas, rentas de carros, consultores, gimnasios, salones de belleza.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Problema: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Invisibilidad en Google, citas manuales, doble reservación.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Posicionamiento: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('"Es como tener una recepcionista digital 24/7."')
run.italic = True
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('IA: ')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = ORANGE
run = p.add_run('Chatbot FAQ, reservas inteligentes, SEO automático para Google.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run('💰 Precios: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Chivo $99 inst./$25/mes · Pro $149/$49/mes · Premium $299/$89/mes')
run.font.size = Pt(8)
run.font.color.rgb = MEDIUM_BLUE

# --- PRODUCT 3: VACATION RENTALS ---
add_shaded_header(doc, '3️⃣  VACATION RENTALS — Plataforma de Reservas Directas')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Ideal para: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Propietarios de Airbnb, casas de playa, cabañas, eco-lodges.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Problema: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Comisiones altas de Airbnb (15%), falta de control y marca propia.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Posicionamiento: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('"Recupere el porcentaje que hoy le paga a Airbnb."')
run.italic = True
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('IA: ')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = ORANGE
run = p.add_run('Precios Dinámicos IA (ajusta tarifas según demanda del mercado).')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run('💰 Precios: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Modelo A $250–350 inst./8–10% com. · Modelo B $400–600 inst./18–20% com.')
run.font.size = Pt(8)
run.font.color.rgb = MEDIUM_BLUE

# --- PRODUCT 4: AUTOMATION ---
add_shaded_header(doc, '4️⃣  AUTOMATIZACIÓN DE PROCESOS')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Ideal para: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Firmas contables, oficinas legales, RRHH, aseguradoras, bienes raíces.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Problema: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Tareas repetitivas, carga administrativa, reportes manuales.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Posicionamiento: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('"Elimine trabajo repetitivo y libere tiempo para crecer."')
run.italic = True
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('IA: ')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = ORANGE
run = p.add_run('Flujos adaptativos, detección de errores en tiempo real.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run('💰 Precios: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Básica desde $200/mes · Pro desde $400/mes · Empresarial precio personalizado')
run.font.size = Pt(8)
run.font.color.rgb = MEDIUM_BLUE

# --- PRODUCT 5: CHATBOT ---
add_shaded_header(doc, '5️⃣  AGENTE CHATBOT CON IA')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Ideal para: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Clínicas, abogados, gimnasios, concesionarios, restaurantes, inmobiliarias.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Problema: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Mensajes fuera de horario, preguntas repetitivas, pérdida de prospectos.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('Posicionamiento: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('"Un empleado virtual que trabaja 24/7 sin descanso."')
run.italic = True
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('IA: ')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = ORANGE
run = p.add_run('Conversación natural español/inglés, aprendizaje continuo, califica prospectos.')
run.font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run('💰 Precios: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Básico desde $150/mes · Pro desde $300/mes · Personalizado cotización')
run.font.size = Pt(8)
run.font.color.rgb = MEDIUM_BLUE

# --- DECISION TREE ---
add_shaded_header(doc, '🌳  ÁRBOL DE DECISIÓN RÁPIDO')

decisions = [
    ('¿Es propietario de inmueble?', '→ Vacation Rentals'),
    ('¿Vende productos o comida?', '→ AlToque App'),
    ('¿Es negocio de servicios?', '→ Desarrollo Web (WaaS)'),
    ('¿Tiene procesos repetitivos?', '→ + Automatización'),
    ('¿Recibe muchas consultas?', '→ + Agente Chatbot IA'),
]

for question, answer in decisions:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(question + '  ')
    run.bold = True
    run.font.size = Pt(8)
    run = p.add_run(answer)
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_BLUE

# --- STACK SELLING ---
add_shaded_header(doc, '📦  VENTA EN PAQUETES (STACK SELLING)')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Regla clave: ')
run.bold = True
run.font.size = Pt(8)
run = p.add_run('Industria → Identificar Dolor → Emparejar Producto → Agregar Paquete')
run.font.size = Pt(8)
run.italic = True

stacks = [
    ['Restaurante Completo', 'AlToque + Web + Chatbot IA', 'Google + pedidos + atención 24/7'],
    ['Propietario de Inmueble', 'Vacation Rentals + Web + Automatización', 'Vitrina + reservas + seguimientos auto'],
    ['Negocio de Servicios', 'Web Pro + Chatbot IA + Automatización', 'Reservas + atención auto + procesos'],
    ['Emprendedor Digital', 'AlToque + Chatbot IA', 'Catálogo + respuesta instantánea'],
]

make_table(doc, ['Paquete', 'Incluye', 'Por Qué Funciona'], stacks, [4.5, 6.5, 6.5])

# Pitch line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('💡 "Con solo la app vas a crecer. Pero si además tienes chatbot y web — es como tener 3 empleados extra. Y te lo hago todo hoy a precio especial."')
run.italic = True
run.font.size = Pt(7.5)
run.font.color.rgb = DARK_GRAY

# ============================================================
# PAGE BREAK
# ============================================================
doc.add_page_break()

# ============================================================
# PAGE 2 — SALES TOOLS
# ============================================================

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2.paragraph_format.space_after = Pt(6)
run = title2.add_run('TODO DIRECTO — HERRAMIENTAS DE VENTA')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = DARK_BLUE

# --- SPIN QUESTIONS ---
add_shaded_header(doc, '🔍  LAS 5 PREGUNTAS DE DESCUBRIMIENTO (SPIN)')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('En Todo Directo no "empujamos" productos. Diagnosticamos problemas. Escucha más de lo que hablas — el 70% de la conversación debe ser el cliente.')
run.font.size = Pt(7.5)
run.italic = True
run.font.color.rgb = DARK_GRAY

spin_questions = [
    ('1. Situación: ', '"¿Cómo maneja actualmente sus pedidos/reservas/consultas de clientes?"'),
    ('2. Problema: ', '"¿Cuántos clientes pierde en una semana porque está muy ocupado para responder rápido?"'),
    ('3. Implicación: ', '"Si recuperara solo 5 ventas perdidas por semana, ¿cuánto dinero extra generaría al mes?"'),
    ('4. Necesidad: ', '"¿Qué significaría para su negocio si los clientes pudieran hacer pedidos 24/7, incluso cuando está cerrado?"'),
    ('5. Calificación: ', 'Varía por tipo → Propietario: "¿Vive cerca?" / Restaurante: "¿Cuántos productos?" / Servicios: "¿Lo buscan en Google?"'),
]

for label, question in spin_questions:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_BLUE
    run = p.add_run(question)
    run.font.size = Pt(8)

# --- OBJECTION HANDLING ---
add_shaded_header(doc, '🛡️  ELIMINADORES DE OBJECIONES')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Cada objeción es una señal de interés. Un cliente que objeta está involucrado.')
run.font.size = Pt(7.5)
run.italic = True
run.font.color.rgb = DARK_GRAY

objections = [
    ['"No tengo dinero"', 'Por eso necesita esto — para GANAR más. Si esperamos 2 meses, ¿cuántos clientes pierde? Se paga solo en la primera semana.'],
    ['"Necesito pensarlo"', '¿Qué parte específicamente? ¿Precio o funcionalidad? (Identifica la duda real.)'],
    ['"Ya tengo sitio web"', '¿Le trae clientes nuevos cada semana? El problema no es tener web, es tener un SISTEMA completo.'],
    ['"Mis clientes no usan tecnología"', 'Pero sus hijos sí. Y los turistas definitivamente sí. Sin esto, solo es visible para el 20% del mercado.'],
    ['"No confío en la IA"', 'Usted siempre tiene el control. El chatbot maneja lo repetitivo — usted maneja lo importante.'],
    ['"Es muy caro"', '¿Cuánto cuesta un empleado de tiempo completo? Este trabaja 24/7 por una fracción de ese costo.'],
]

make_table(doc, ['Objeción', 'Tu Respuesta'], objections, [4.5, 13.5])

# --- COMMISSION RULES ---
add_shaded_header(doc, '💰  LAS 4 REGLAS DE COMISIONES')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Ganas de dos formas: Comisión por cierre (inmediato) + Residual mensual (mientras el cliente siga activo).')
run.font.size = Pt(7.5)
run.italic = True
run.font.color.rgb = DARK_GRAY

commission_rules = [
    ['Regla 40/10', 'Desarrollo Web', '40% de instalación + 10% mensual', 'Pro: $59.60 + $4.90/mes'],
    ['Regla 50/15', 'Automatización y Chatbot', '50% del 1er mes + 15% mensual', 'Chat Básico: $75 + $22.50/mes'],
    ['Regla 150%', 'AlToque', '1.5× la mensualidad (1er mes) + 10% mensual', 'Pro: $22.49 + $1.50/mes'],
    ['Regla Fija', 'Vacation Rentals', '$300 (Modelo A) o $600 (Modelo B)', 'Siempre fijo, sin residual'],
]

make_table(doc, ['Regla', 'Producto', 'Cómo Funciona', 'Ejemplo'], commission_rules, [3, 4, 5.5, 5.5])

# Residual callout
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(2)
run = p.add_run('⚠️ POTENCIAL REAL: ')
run.bold = True
run.font.size = Pt(8)
run.font.color.rgb = ORANGE
run = p.add_run('Un rep con 100 clientes activos genera ~$1,330/mes en residual puro — sin cerrar nada nuevo. Construye tu portafolio desde el día 1 y protégelo.')
run.font.size = Pt(8)

# Volume bonuses mini-table
add_section_header(doc, 'Bonos por Volumen')

bonuses = [
    ['Semanal', '5 tratos/semana', '+$50'],
    ['Mensual', '20 tratos/mes', '+$200'],
    ['Trimestral', '60 tratos en 3 meses', '+$500'],
]

make_table(doc, ['Bono', 'Objetivo', 'Premio'], bonuses, [4, 7, 4])

# --- CLOSING PHRASES ---
add_shaded_header(doc, '🎯  FRASES DE CIERRE')

closing_phrases = [
    ('Descuento anual: ', '"Si paga el año hoy, le regalo 2 meses gratis."'),
    ('Urgencia: ', '"¿Cuándo quiere que esté listo? Puedo tenerlo la próxima semana."'),
    ('ROI: ', '"Si recupera solo 3 ventas perdidas por semana, esto se paga solo."'),
    ('Visión: ', '"Imagínese abrir su WhatsApp mañana y ver que su chatbot ya cerró 5 pedidos mientras usted dormía."'),
]

for label, phrase in closing_phrases:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_BLUE
    run = p.add_run(phrase)
    run.font.size = Pt(8)

# --- COLD OPEN SCRIPTS ---
add_shaded_header(doc, '🚪  APERTURA EN FRÍO (3 Versiones)')

cold_opens = [
    ('Local: ', '"Buenos días, ¿es usted el dueño? Soy [nombre] de Todo Directo. Ayudamos a negocios como el suyo a no perder clientes por estar muy ocupados. ¿Tiene un minuto?"'),
    ('Propietario: ', '"Hola, vi que tiene esta propiedad en Airbnb. ¿Sabía que podría ganar un 10% más por noche usando una plataforma local sin las comisiones altas?"'),
    ('Automatización: ', '"Hola, soy [nombre] de Todo Directo. Trabajamos con negocios para recuperar las horas que pierden respondiendo mensajes repetitivos con IA."'),
]

for label, script in cold_opens:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = MEDIUM_BLUE
    run = p.add_run(script)
    run.font.size = Pt(8)
    run.italic = True

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Todo Directo  |  www.tododirecto.com  |  info@tododirecto.com  |  El Salvador  |  v2.0 Feb 2026')
run.font.size = Pt(7)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================================================
# SAVE
# ============================================================
output_path = 'Todo_Directo_2_Paginas_Sales_Cheat_Sheet_ES_v2.docx'
doc.save(output_path)
print(f"✅ Saved: {output_path}")
