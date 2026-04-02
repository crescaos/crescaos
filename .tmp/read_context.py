import pdfplumber
from pypdf import PdfReader
import openpyxl

# 1. Read all context documents
print("="*80)
print("PRICING STRATEGY (docx via pypdf won't work - use python-docx)")
print("="*80)

try:
    import docx
    doc = docx.Document('Todo_Directo_Pricing_Strategy.docx')
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
except:
    print("python-docx not installed, will install")

print("\n" + "="*80)
print("SOP v2")
print("="*80)
try:
    doc2 = docx.Document('Todo_Directo_SOP_v2.docx')
    for para in doc2.paragraphs:
        if para.text.strip():
            print(para.text)
except:
    print("python-docx not installed")

print("\n" + "="*80)
print("INTAKE FORM")
print("="*80)
try:
    doc3 = docx.Document('Todo_Directo_Intake_Form.docx')
    for para in doc3.paragraphs:
        if para.text.strip():
            print(para.text)
except:
    print("python-docx not installed")
