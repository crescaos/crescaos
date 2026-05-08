"""Generate CRESCA_Contrato_de_Servicios.docx from the markdown source."""
import os, sys, re

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT

BASE = r"c:\Users\12132\Desktop\DigiFacil\DigiFacil"
MD_PATH = os.path.join(BASE, "CRESCA_Contrato_de_Servicios.md")
OUT_PATH = os.path.join(BASE, "CRESCA_Contrato_de_Servicios.docx")

with open(MD_PATH, "r", encoding="utf-8") as f:
    lines = f.readlines()

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Style setup ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    sname = f'Heading {level}'
    s = doc.styles[sname]
    s.font.name = 'Times New Roman'
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        s.font.size = Pt(16)
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after = Pt(12)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        s.font.size = Pt(14)
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(8)
    else:
        s.font.size = Pt(12)
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(6)


def add_formatted_paragraph(doc, text, style='Normal', bold=False, alignment=None, indent=False):
    """Add a paragraph with inline bold formatting."""
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    if alignment:
        p.paragraph_format.alignment = alignment

    # Split by **...**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def parse_table(table_lines):
    """Parse markdown table lines into rows of cells."""
    rows = []
    for line in table_lines:
        line = line.strip().strip('|')
        if re.match(r'^[\s\-\|]+$', line):
            continue
        cells = [c.strip() for c in line.split('|')]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold
            parts = re.split(r'(\*\*.*?\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
            if i == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()  # spacing


# --- Process markdown lines ---
i = 0
skip_note = False
while i < len(lines):
    line = lines[i].rstrip('\n')
    stripped = line.strip()

    # Skip the Word formatting note block
    if stripped.startswith('> **NOTA PARA MICROSOFT WORD'):
        i += 1
        continue

    # Skip empty lines
    if not stripped:
        i += 1
        continue

    # Skip horizontal rules
    if stripped == '---':
        i += 1
        continue

    # Skip &nbsp;
    if stripped == '&nbsp;':
        i += 1
        continue

    # Tables
    if '|' in stripped and stripped.startswith('|'):
        table_lines = []
        while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        rows = parse_table(table_lines)
        add_table(doc, rows)
        continue

    # H1
    if stripped.startswith('# ') and not stripped.startswith('## '):
        heading_text = stripped[2:].strip()
        heading_text = re.sub(r'\*\*', '', heading_text)
        doc.add_heading(heading_text, level=1)
        i += 1
        continue

    # H2
    if stripped.startswith('## '):
        heading_text = stripped[3:].strip()
        heading_text = re.sub(r'\*\*', '', heading_text)
        doc.add_heading(heading_text, level=2)
        i += 1
        continue

    # H3
    if stripped.startswith('### '):
        heading_text = stripped[4:].strip()
        heading_text = re.sub(r'\*\*', '', heading_text)
        doc.add_heading(heading_text, level=3)
        i += 1
        continue

    # Blockquote / indented content
    if stripped.startswith('> '):
        text = stripped[2:].strip()
        # Checkbox lines
        if '☐' in text:
            add_formatted_paragraph(doc, text, indent=True)
        else:
            add_formatted_paragraph(doc, text, indent=True)
        i += 1
        continue

    # Signature lines
    if stripped.startswith('___'):
        p = doc.add_paragraph('_' * 50)
        p.paragraph_format.space_before = Pt(24)
        i += 1
        continue

    # Normal paragraph
    # Clean up markdown artifacts
    text = stripped
    add_formatted_paragraph(doc, text)
    i += 1

# --- Save ---
doc.save(OUT_PATH)
print(f"Document saved to: {OUT_PATH}")
print("Done!")
