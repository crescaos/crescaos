# Part 1: Document setup + Sections 1-4
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import pickle, os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for lvl, sz, clr in [('Heading 1', 16, '1B4F72'), ('Heading 2', 13, '2E86C1'), ('Heading 3', 11, '2874A6')]:
    s = doc.styles[lvl]
    s.font.name = 'Calibri'
    s.font.size = Pt(sz)
    s.font.color.rgb = RGBColor.from_string(clr)
    s.font.bold = True

PURPLE = RGBColor.from_string('6C3483')
DARK = RGBColor.from_string('1B2631')
HEADER_BG = '2E86C1'

def set_cell_bg(cell, color):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(s)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255,255,255)
                r.font.size = Pt(10)
        set_cell_bg(c, HEADER_BG)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Inches(w)
    doc.add_paragraph()
    return t

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level*0.25)
    return p

# ========== TITLE PAGE ==========
for _ in range(6): doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('TODO DIRECTO')
r.font.size = Pt(28); r.bold = True; r.font.color.rgb = PURPLE

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Manual de Operaciones Internas')
r2.font.size = Pt(18); r2.font.color.rgb = DARK

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Procedimientos de Onboarding, Comunicación, Calidad y KPIs')
r3.font.size = Pt(12); r3.font.color.rgb = RGBColor.from_string('566573')

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run('Versión 1.0  |  Marzo 2026  |  Confidencial – Solo Uso Interno')
r4.font.size = Pt(10); r4.font.color.rgb = RGBColor.from_string('808B96')

doc.add_page_break()

# ========== TABLE OF CONTENTS ==========
doc.add_heading('TABLA DE CONTENIDOS', level=1)
toc_items = [
    '1. Propósito y Alcance',
    '2. Roles y Responsabilidades',
    '3. Pipeline de Onboarding de Clientes (7 Etapas)',
    '4. Checklists de Onboarding por Producto',
    '5. Reglas de Comunicación Interna',
    '6. Plantillas de Comunicación con el Cliente',
    '7. Reglas de Planificación de Capacidad',
    '8. Protocolo de Escalamiento',
    '9. Estrategia de Prevención de Churn',
    '10. Checklist de Control de Calidad',
    '11. Estándares de Documentación y Archivos',
    '12. KPIs Operacionales',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
doc.add_page_break()

# ========== SECTION 1 ==========
doc.add_heading('SECCIÓN 1: PROPÓSITO Y ALCANCE', level=1)
doc.add_paragraph(
    'Este documento define los procedimientos operacionales internos de Todo Directo — la empresa de '
    'transformación digital para negocios en El Salvador. Cubre todos los procesos desde que un '
    'representante comercial cierra una venta hasta que el cliente está activo, y los procedimientos '
    'continuos de soporte, calidad y retención.'
)
doc.add_heading('¿Quién Debe Usar Este Manual?', level=2)
add_bullet(doc, 'Equipo de Operaciones — responsable de implementar y entregar los servicios contratados.')
add_bullet(doc, 'Soporte Técnico — responsable de resolver problemas y mantener la calidad del servicio.')
add_bullet(doc, 'Gerente de Operaciones — responsable de supervisar el pipeline, SLAs y KPIs.')
add_bullet(doc, 'Fundador/CEO — para referencia y toma de decisiones de escalamiento.')

doc.add_heading('¿Qué NO Cubre Este Manual?', level=2)
doc.add_paragraph(
    'Este manual NO cubre el proceso de ventas en campo. Para scripts de venta, manejo de objeciones, '
    'comisiones y técnicas de cierre, consultar el SOP de Ventas (Todo_Directo_SOP_v2.docx).'
)

doc.add_heading('Control de Versiones', level=2)
add_table(doc, ['Versión', 'Fecha', 'Cambios', 'Autor'],
    [['1.0', 'Marzo 2026', 'Documento inicial', '[PLACEHOLDER]']],
    [1.0, 1.5, 3.0, 1.5])

# ========== SECTION 2 ==========
doc.add_heading('SECCIÓN 2: ROLES Y RESPONSABILIDADES', level=1)
doc.add_paragraph(
    'Todo Directo opera con un equipo lean. Cada rol tiene responsabilidades claras para evitar '
    'duplicación de esfuerzos y asegurar que ningún cliente quede sin atención.'
)

doc.add_heading('Organigrama Funcional', level=2)
add_table(doc, ['Rol', 'Responsabilidades Principales', 'Reporta A'],
    [
        ['Fundador / CEO', 'Estrategia, decisiones de precios, escalamientos Nivel 3, alianzas', '—'],
        ['Gerente de Operaciones', 'Pipeline de onboarding, SLAs, KPIs, coordinación inter-equipos', 'CEO'],
        ['Equipo Técnico (Desarrollo)', 'Implementación de sitios web, apps, chatbots, automatizaciones', 'Gerente Ops'],
        ['Soporte / QA', 'Control de calidad pre-lanzamiento, soporte post-lanzamiento, bugs', 'Gerente Ops'],
        ['Representantes Comerciales', 'Captura de leads, cierre de ventas, entrega de formulario de intake', 'CEO'],
    ],
    [2.0, 3.5, 1.5])

doc.add_heading('Matriz RACI — Procesos Clave', level=2)
doc.add_paragraph('R = Responsable (ejecuta) | A = Aprobador | C = Consultado | I = Informado')
add_table(doc, ['Proceso', 'Rep Comercial', 'Gerente Ops', 'Equipo Técnico', 'Soporte/QA', 'CEO'],
    [
        ['Cierre de venta', 'R', 'I', 'I', '—', 'I'],
        ['Formulario de intake', 'R', 'A', 'I', '—', '—'],
        ['Kickoff técnico', 'I', 'R', 'R', '—', '—'],
        ['Implementación', '—', 'C', 'R', '—', '—'],
        ['QA pre-lanzamiento', '—', 'A', 'C', 'R', '—'],
        ['Go-live', '—', 'R', 'R', 'R', 'I'],
        ['Soporte post-launch', '—', 'I', 'C', 'R', '—'],
        ['Escalamiento Nivel 3', '—', 'R', 'C', 'I', 'A'],
        ['Decisión de precios', '—', 'C', '—', '—', 'R/A'],
        ['Retención / anti-churn', 'C', 'R', 'C', 'R', 'I'],
    ],
    [2.0, 1.0, 1.0, 1.0, 1.0, 0.8])

# ========== SECTION 3 ==========
doc.add_heading('SECCIÓN 3: PIPELINE DE ONBOARDING DE CLIENTES', level=1)
doc.add_paragraph(
    'Todo cliente nuevo pasa por un pipeline de 7 etapas. Cada etapa tiene un SLA (tiempo máximo '
    'permitido) y un responsable claro. El Gerente de Operaciones supervisa que ningún cliente '
    'se quede estancado en una etapa.'
)

doc.add_heading('Vista General del Pipeline', level=2)
add_table(doc, ['Etapa', 'Nombre', 'SLA', 'Responsable', 'Entregable'],
    [
        ['1', 'Captura y Asignación del Lead', '≤ 2 horas', 'Gerente Ops', 'Proyecto creado en sistema de tracking'],
        ['2', 'Formulario de Intake Completo', '≤ 24 horas', 'Rep Comercial', 'Formulario completo + archivos del cliente'],
        ['3', 'Pago y Contrato', 'Mismo día del cierre', 'Rep Comercial', 'Pago confirmado + contrato firmado'],
        ['4', 'Kickoff Técnico', '≤ 48 horas post-pago', 'Gerente Ops + Técnico', 'Reunión de kickoff completada'],
        ['5', 'Implementación y Construcción', 'Ver tabla por producto', 'Equipo Técnico', 'Producto construido y funcional'],
        ['6', 'Revisión del Cliente', '≤ 48 horas', 'Soporte/QA', 'Aprobación del cliente o lista de cambios'],
        ['7', 'Go-Live y Activación', '≤ 24 horas post-aprobación', 'Equipo Técnico', 'Producto en producción + cliente notificado'],
    ],
    [0.5, 1.8, 1.5, 1.5, 2.0])

# Detailed stages
stages = [
    ('Etapa 1: Captura y Asignación del Lead', [
        'El representante comercial cierra la venta y envía el reporte al equipo de Todo Directo.',
        'El Gerente de Operaciones recibe la notificación y crea el proyecto en el sistema de tracking.',
        'Se asigna un técnico responsable según el tipo de producto contratado.',
        'Se envía confirmación al representante y al cliente.',
    ], 'SLA: Máximo 2 horas desde que el representante reporta el cierre.'),
    ('Etapa 2: Formulario de Intake Completo', [
        'El representante debe haber completado el Formulario de Orden de Servicio con el cliente en el momento del cierre.',
        'El formulario se escanea y envía a info@tododirecto.com el mismo día.',
        'Gerente Ops verifica que todos los campos requeridos estén completos.',
        'Si faltan campos o archivos, se contacta al cliente dentro de 4 horas para solicitar lo faltante.',
    ], 'SLA: El formulario debe estar completo y verificado dentro de 24 horas del cierre.'),
    ('Etapa 3: Pago y Contrato', [
        'La cuota de instalación debe cobrarse al momento del cierre (Efectivo, Transferencia, Chivo Wallet, Tarjeta).',
        'El representante confirma el pago y envía comprobante.',
        'Se genera y firma el contrato de servicios.',
        'El pago se registra en el Dashboard Financiero.',
    ], 'SLA: Mismo día del cierre. No se inicia implementación sin pago confirmado.'),
    ('Etapa 4: Kickoff Técnico', [
        'Reunión interna (15–30 min) entre Gerente Ops y técnico asignado.',
        'Revisar: producto contratado, plan seleccionado, información del intake, archivos del cliente.',
        'Definir timeline específico según complejidad.',
        'Enviar mensaje de bienvenida al cliente con timeline estimado.',
    ], 'SLA: Máximo 48 horas después del pago confirmado.'),
    ('Etapa 5: Implementación y Construcción', [
        'El equipo técnico construye el producto según las especificaciones del intake.',
        'Se siguen los checklists por producto (ver Sección 4).',
        'Actualizaciones de progreso al Gerente Ops cada 48 horas en proyectos de más de 5 días.',
        'Si se necesita información adicional del cliente, se solicita inmediatamente — no esperar.',
    ], None),
    ('Etapa 6: Revisión del Cliente', [
        'QA interno completa el checklist de calidad (ver Sección 10) antes de mostrar al cliente.',
        'Se agenda una llamada o visita con el cliente para revisar el producto.',
        'Se documentan cambios solicitados (máximo 2 rondas de revisión incluidas).',
        'Cambios adicionales se cotizan por separado.',
    ], 'SLA: El cliente tiene 48 horas para dar feedback. Si no responde, se envía recordatorio.'),
    ('Etapa 7: Go-Live y Activación', [
        'Producto se publica/activa en producción.',
        'Se envía notificación de go-live al cliente con instrucciones de uso.',
        'Se activa la suscripción mensual en el sistema de cobros.',
        'Se agenda el primer check-in post-lanzamiento (a las 2 semanas).',
        'Se notifica al representante comercial que su cliente está activo.',
    ], 'SLA: Máximo 24 horas después de la aprobación del cliente.'),
]

for title, steps_list, sla in stages:
    doc.add_heading(title, level=2)
    for s in steps_list:
        add_bullet(doc, s)
    if sla:
        p = doc.add_paragraph()
        r = p.add_run('⏱ ' + sla)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string('C0392B')
    doc.add_paragraph()

doc.add_heading('Tiempos de Implementación por Producto (Etapa 5)', level=2)
add_table(doc, ['Producto', 'Plazo Estimado', 'Qué Se Necesita del Cliente'],
    [
        ['AlToque App', '24–48 horas', 'Fotos de productos, precios, WhatsApp'],
        ['Desarrollo Web – Plan Chivo', '5–7 días hábiles', 'Logo, fotos, descripción del negocio'],
        ['Desarrollo Web – Plan Pro/Premium', '7–10 días hábiles', 'Todo anterior + horarios, servicios'],
        ['PR Vacation Rentals', '5–7 días hábiles', 'Fotos propiedad, precios, disponibilidad'],
        ['Automatización de Procesos', '7–14 días hábiles', 'Flujos actuales, integraciones requeridas'],
        ['Chatbot IA', '5–10 días hábiles', 'Preguntas frecuentes, tono de marca'],
    ],
    [2.0, 1.5, 3.5])

# ========== SECTION 4 ==========
doc.add_heading('SECCIÓN 4: CHECKLISTS DE ONBOARDING POR PRODUCTO', level=1)
doc.add_paragraph(
    'Cada producto tiene un checklist específico que el equipo técnico debe completar durante la '
    'implementación. Estos checklists aseguran consistencia y calidad en cada entrega.'
)

products = [
    ('4.1 AlToque App', [
        ('Configuración Inicial', [
            'Crear cuenta del negocio en plataforma AlToque',
            'Configurar nombre del negocio, logo, WhatsApp de contacto',
            'Seleccionar plan (Chivo/Pro/Premium) según contrato',
            'Configurar zona de entrega (si aplica)',
        ]),
        ('Catálogo de Productos', [
            'Subir todas las fotos de productos proporcionadas por el cliente',
            'Ingresar nombres, descripciones y precios de cada producto',
            'Organizar productos por categorías',
            'Verificar que cada producto tenga foto, nombre y precio correcto',
        ]),
        ('Pruebas', [
            'Realizar pedido de prueba completo',
            'Verificar que la notificación llega al WhatsApp del cliente',
            'Probar en móvil y desktop',
            'Verificar analíticas IA funcionando (si es Pro/Premium)',
        ]),
    ]),
    ('4.2 Desarrollo Web / App (WaaS)', [
        ('Configuración', [
            'Registrar dominio o configurar subdominio',
            'Configurar hosting y SSL',
            'Aplicar logo y colores de marca del cliente',
        ]),
        ('Contenido', [
            'Crear páginas según plan: inicio, servicios, contacto',
            'Integrar sistema de reservas/citas (Plan Pro+)',
            'Configurar formulario de contacto',
            'Optimizar SEO básico (títulos, meta descriptions)',
            'Configurar chatbot IA para FAQ (si aplica)',
        ]),
        ('Pruebas', [
            'Verificar responsive (móvil, tablet, desktop)',
            'Probar formulario de contacto — verificar que llega al email del cliente',
            'Probar sistema de reservas — crear cita de prueba',
            'Verificar velocidad de carga < 3 segundos',
            'Verificar que aparece en búsqueda de Google (post-indexación)',
        ]),
    ]),
    ('4.3 PR Vacation Rentals', [
        ('Configuración', [
            'Crear perfil de propiedad en plataforma',
            'Subir fotos profesionales de la propiedad',
            'Configurar amenidades, reglas de la casa, capacidad',
            'Definir modelo contratado (A: Direct Booking o B: Full Management)',
        ]),
        ('Calendario y Precios', [
            'Configurar calendario de disponibilidad',
            'Establecer precios por temporada (alta, media, baja)',
            'Configurar precios dinámicos IA (si aplica)',
            'Sincronizar con Airbnb/Booking si el cliente lo requiere',
        ]),
        ('Pruebas', [
            'Hacer una reserva de prueba completa',
            'Verificar flujo de confirmación automática',
            'Probar panel de control del propietario',
            'Verificar que el sistema de cobro funciona',
        ]),
    ]),
    ('4.4 Automatización de Procesos', [
        ('Descubrimiento', [
            'Documentar proceso actual del cliente paso a paso',
            'Identificar los 1–3 flujos más críticos a automatizar primero',
            'Mapear integraciones necesarias (WhatsApp, Gmail, Google Sheets, etc.)',
        ]),
        ('Implementación', [
            'Construir flujos automatizados según especificación',
            'Conectar todas las integraciones requeridas',
            'Configurar triggers (eventos que disparan la automatización)',
            'Agregar manejo de errores y notificaciones de fallo',
        ]),
        ('Pruebas', [
            'Ejecutar cada flujo con datos de prueba',
            'Verificar que los triggers se disparan correctamente',
            'Probar escenarios de error — verificar que se notifican',
            'Confirmar con el cliente que el resultado es el esperado',
        ]),
    ]),
    ('4.5 Agentes Chatbot con IA', [
        ('Configuración', [
            'Definir canal(es): WhatsApp Business, Instagram, sitio web, Facebook Messenger',
            'Recopilar y cargar preguntas frecuentes del cliente (5–10 mínimo)',
            'Configurar tono de comunicación (formal, casual, directo, entusiasta)',
            'Definir horario de atención humana vs. chatbot 24/7',
        ]),
        ('Entrenamiento IA', [
            'Entrenar el modelo con información específica del negocio',
            'Configurar respuestas para las FAQ más comunes',
            'Definir reglas de escalamiento a humano',
            'Configurar idiomas (español, inglés si aplica)',
        ]),
        ('Pruebas', [
            'Hacer 10+ preguntas de prueba en cada canal configurado',
            'Verificar que escala a humano cuando corresponde',
            'Probar fuera de horario laboral',
            'Verificar tono y calidad de las respuestas',
            'Confirmar que el cliente puede ver las conversaciones',
        ]),
    ]),
]

for title, phases in products:
    doc.add_heading(title, level=2)
    for phase_name, items in phases:
        doc.add_heading(phase_name, level=3)
        for item in items:
            add_bullet(doc, '☐  ' + item)
    doc.add_paragraph()

# Save intermediate
doc.save(r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\.tmp\sop_partial.docx')
print('Part 1 done - Sections 1-4 saved')
