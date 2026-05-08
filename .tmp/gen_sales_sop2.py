"""Generate Sales SOP Part 2 (Sections 8-13) and merge with Part 1."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document(r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\.tmp\sales_sop_p1.docx')

def ap(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p

def at(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c = t.cell(0,j); c.text = ''; run = c.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.name = 'Calibri'
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = t.cell(i+1,j); c.text = ''; run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10); run.font.name = 'Calibri'
    doc.add_paragraph(); return t

# === SECTION 8 ===
doc.add_heading('SECCIÓN 8: CRESCA COMO TU DIFERENCIADOR', level=1)
ap('La Promesa de Cresca', bold=True)
ap('"No vendemos páginas web. No vendemos chatbots. No vendemos marketing. Vendemos SISTEMAS que hacen que tu negocio funcione mejor, capture más clientes y crezca de forma estructurada — sin que tengas que contratar más gente."')
ap('')
ap('Frases de posicionamiento para usar con clientes:', bold=True)
ap('• "Somos una empresa de sistemas de crecimiento empresarial."')
ap('• "Construimos la infraestructura digital de tu negocio."')
ap('• "Empiece pequeño y escale cuando crezca — desde $9.99 al mes."')
ap('• "Tu negocio funciona 24/7, incluso cuando tú no estás."')
ap('• "No es un gasto — es una inversión que se paga sola."')

ap('')
ap('Lo que nos diferencia de la competencia:', bold=True)
at(['Competencia', 'Cresca'], [
    ['Vende servicios sueltos', 'Vende sistemas integrados en 3 niveles'],
    ['Contrato por proyecto', 'Relación a largo plazo con soporte continuo'],
    ['Te deja solo después de entregar', 'Monitoreo, mantenimiento y optimización continua'],
    ['Un precio único que no escala', 'Escalera de precios: empieza en $9.99, escala a medida'],
    ['Marketing genérico', 'Infraestructura + automatización personalizada'],
])

# === SECTION 9 ===
doc.add_heading('SECCIÓN 9: GUÍA DE MANEJO DE OBJECIONES', level=1)

doc.add_heading('"No tengo dinero ahora mismo"', level=2)
ap('Respuesta:', bold=True)
ap('"Entiendo. Por eso creamos STARTER — por solo $9.99 al mes puede tener un workflow automatizado que le ayude a no perder clientes. Es menos de lo que cuesta un café al día."')
ap('"Además, ofrecemos un Plan Flex: puede dividir el setup en 50% hoy y 50% en 30 días."')
ap('Si definitivamente no puede pagar ni $9.99 → No es nuestro cliente. No insistir.')

doc.add_heading('"Necesito pensarlo"', level=2)
ap('Respuesta:', bold=True)
ap('"Entiendo perfectamente. ¿Puedo preguntar qué es lo que necesitas pensar? Quiero asegurarme de haber respondido todas tus dudas."')
ap('Si es genuino: "Te envío un resumen por WhatsApp para que lo revises. ¿Puedo llamarte el [día]?"')
ap('Si es excusa: "Cada día sin un sistema es un día que pierdes clientes. ¿Qué necesitaría ver para tomar la decisión hoy?"')

doc.add_heading('"Ya tengo un sitio web / Facebook"', level=2)
ap('Respuesta:', bold=True)
ap('"Genial que ya tengas algo. La pregunta es: ¿te está generando clientes de forma automática? ¿Tienes workflows que hagan seguimiento? Cresca no reemplaza lo que tienes — lo potencia con automatización y workflows que trabajan solos."')
ap('"Si ya tiene web, el valor de Cresca está en los WORKFLOWS y AUTOMATIZACIÓN, no en otra página más."')

doc.add_heading('"Mis clientes no usan tecnología"', level=2)
ap('Respuesta:', bold=True)
ap('"Tus clientes usan WhatsApp, ¿verdad? Usan Google. La tecnología que implementamos no requiere que tus clientes sean expertos — solo necesitan encontrarte y contactarte, y el sistema hace el resto."')

doc.add_heading('"No confío en la inteligencia artificial"', level=2)
ap('Respuesta:', bold=True)
ap('"La automatización que usamos no reemplaza a las personas — las ayuda. Es como tener un asistente que trabaja 24/7 y le pasa las conversaciones importantes a ti. Tú siempre tienes el control."')

doc.add_heading('"Ya tengo alguien que me ayuda"', level=2)
ap('Respuesta:', bold=True)
ap('"Perfecto, ¿y esa persona te da todo integrado? ¿Workflows automáticos + CRM + presencia web + soporte continuo por un precio mensual? Cresca no es una persona — es un sistema completo que trabaja las 24 horas."')

doc.add_heading('"Es muy caro"', level=2)
ap('Respuesta:', bold=True)
ap('"STARTER empieza en $9.99 al mes — menos que un almuerzo. Y si necesitas más infraestructura, GROWTH empieza en $50. ¿Cuánto cuesta un empleado de tiempo completo al mes? Este sistema hace el trabajo por una fracción de ese costo."')

doc.add_heading('Guía de Retención: Cuándo un Cliente Quiere Cancelar', level=2)
ap('1. NO ofrezcas descuento inmediatamente.', bold=True)
ap('2. Pregunta la razón real de la cancelación.')
ap('3. Si es un problema resoluble, propón un plan de acción.')
ap('4. Si es presupuestario, ofrece bajar a un tier inferior (ej: Growth → Starter Pro).')
ap('5. Registra la razón en el CRM para análisis futuro.')

# === SECTION 10 ===
doc.add_heading('SECCIÓN 10: ESTRUCTURA DE COMPENSACIÓN — REP FUNDADOR ES', level=1)

doc.add_heading('Compensación del Representante Fundador (El Salvador)', level=2)
at(['Concepto', 'Detalle'], [
    ['Salario Base', '$600/mes (garantizado)'],
    ['Comisión de Setup', '20% de los fees de setup cobrados'],
    ['Comisión Residual', '6% del MRR cobrado mensualmente'],
    ['Clawback', 'Si cliente cancela en 60 días por mal fit → 50% de comisión setup devuelta'],
])

doc.add_heading('Reglas de Pago', level=2)
ap('• Comisiones se pagan SOLO sobre dinero cobrado (no facturado)')
ap('• Residual se paga SOLO mientras el cliente esté activo y al corriente')
ap('• No hay comisión sobre costos de terceros (ad spend, fees externos, reembolsos)')
ap('• Ver documento separado: CRESCA_ES_Founding_Rep_Compensation.docx', italic=True)

doc.add_heading('Ejemplo: Ingreso Mensual del Rep Fundador', level=2)
ap('Supuestos: cierra 3 clientes/mes con mix de STARTER y GROWTH')
at(['Mes', 'Clientes Activos', 'Salario Base', 'Setup (20%)', 'Residual (6%)', 'Total Mes'], [
    ['1', '3', '$600', '$60–$100', '$9–$18', '$669–$718'],
    ['3', '9', '$600', '$60–$100', '$27–$54', '$687–$754'],
    ['6', '18', '$600', '$60–$100', '$54–$108', '$714–$808'],
    ['12', '36', '$600', '$60–$100', '$108–$216', '$768–$916'],
])
ap('El residual es acumulativo — crece cada mes que más clientes están activos. 📈', bold=True)

# === SECTION 11 ===
doc.add_heading('SECCIÓN 11: HERRAMIENTAS Y RECURSOS DE VENTAS', level=1)

doc.add_heading('Tu Kit de Ventas', level=2)
at(['Herramienta', 'Uso', 'Dónde Encontrarlo'], [
    ['Pitch Deck', 'Presentación para clientes', 'CRESCA_Pitch_Deck_UPDATED.pptx'],
    ['Sales Cheat Sheet', 'Referencia rápida en campo', 'Sales_Cheat_Sheet_ES_v2.docx'],
    ['Formulario de Intake', 'Recopilar info del prospecto', 'Intake_Form.docx'],
    ['Contrato de Servicios', 'Cierre formal', 'CRESCA_Contrato_de_Servicios.docx'],
    ['Comp Plan (Founding Rep)', 'Referencia de compensación', 'CRESCA_ES_Founding_Rep_Compensation.docx'],
    ['Demo/Prototipos', 'Mostrar lo que Cresca construye', 'Sitios demo + workflows demo'],
])

# === SECTION 12 ===
doc.add_heading('SECCIÓN 12: GUÍA DE CAMPO — FLUJO DE TRABAJO DIARIO', level=1)

doc.add_heading('Planificación del Territorio', level=2)
ap('• Divide tu territorio en zonas de alta y baja densidad.')
ap('• Prioriza zonas con concentración de negocios objetivo.')
ap('• Alterna entre visitas frías y seguimiento de prospectos calientes.')
ap('• Mínimo 10 contactos nuevos por día (frío o tibio).')

doc.add_heading('Estructura del Día', level=2)
at(['Hora', 'Actividad'], [
    ['8:00–8:30', 'Planificación: revisar pipeline, priorizar seguimientos'],
    ['8:30–12:00', 'Prospección activa: visitas frías, auditorías de 15 min, WhatsApp'],
    ['12:00–13:00', 'Almuerzo + documentar contactos en CRM'],
    ['13:00–16:00', 'Reuniones de presentación / demo / cierre'],
    ['16:00–17:00', 'Seguimiento: enviar cotizaciones, responder mensajes'],
    ['17:00–17:30', 'Reporte diario: actualizar CRM, reportar resultados'],
])

# === SECTION 13 ===
doc.add_heading('SECCIÓN 13: TARJETA DE REFERENCIA RÁPIDA', level=1)

doc.add_heading('Auditoría en 15 Minutos — Las 5 Áreas', level=2)
ap('1. Leads — ¿De dónde vienen sus clientes?')
ap('2. Velocidad de respuesta — ¿Cuánto tardan en responder?')
ap('3. Reservas/Ventas — ¿Cómo cierran ventas?')
ap('4. Operaciones — ¿Qué procesos consumen más tiempo?')
ap('5. Retención — ¿Cómo mantienen contacto con clientes?')

doc.add_heading('Árbol de Decisión Rápido', level=2)
ap('No listo digitalmente → STARTER ($9.99–$49.99)', bold=True)
ap('Necesita bookings/leads + automatización → GROWTH ($50–$100)', bold=True)
ap('Dolor operacional + workflows complejos → PRO AUTOMATION (Cotización)', bold=True)
ap('¿Necesita ads? → + Add-On Marketing (solo con Growth o Pro)')
ap('¿Necesita más workflows? → + Bundle de workflows')

doc.add_heading('Frases de Cierre', level=2)
ap('• "¿Procedemos con la configuración esta semana?"')
ap('• "STARTER empieza en $9.99 — menos que un almuerzo al día."')
ap('• "Ofrecemos Plan Flex: 50% del setup hoy, 50% en 30 días."')
ap('• "¿Cuántos clientes más necesitas perder antes de invertir?"')

doc.add_heading('Eliminadores de Objeciones', level=2)
ap('• "No tengo dinero" → "STARTER empieza en $9.99. También hay Plan Flex."')
ap('• "Ya tengo web" → "¿Te genera clientes en automático? El valor está en workflows."')
ap('• "Necesito pensarlo" → "¿Qué información te falta para decidir?"')
ap('• "No ahora" → "¿Cuándo sería buen momento para dejar de perder clientes?"')
ap('• "Muy caro" → "¿Cuánto cuesta NO tener un sistema?"')

ap('')
ap('Cresca — Sistemas de Crecimiento Empresarial', bold=True)
ap('Manual de Ventas — Documento confidencial para uso interno. Versión 3.0 — Marzo 2026', italic=True)

# Save final
OUT = r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\CRESCA_SOP_v2.docx'
doc.save(OUT)
print(f"Saved: {OUT}")
print("Done!")
