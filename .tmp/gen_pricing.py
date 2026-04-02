"""Generate the updated Cresca Pricing Strategy document — 3-Tier Model."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx -q")
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = r"c:\Users\12132\Desktop\DigiFacil\DigiFacil\CRESCA Estrategia de Precio.docx"
doc = Document()

# -- Page setup --
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

# -- Styles --
sn = doc.styles['Normal']
sn.font.name = 'Calibri'
sn.font.size = Pt(11)
sn.paragraph_format.space_after = Pt(6)

for lvl in [1, 2, 3]:
    st = doc.styles[f'Heading {lvl}']
    st.font.name = 'Calibri'
    st.font.bold = True
    st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if lvl == 1:
        st.font.size = Pt(18)
    elif lvl == 2:
        st.font.size = Pt(14)
    else:
        st.font.size = Pt(12)


def add_para(text, style='Normal', bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    doc.add_paragraph()
    return t


# ========== DOCUMENT CONTENT ==========

doc.add_heading('CRESCA — Estrategia de Precios', level=1)
add_para('Sistemas de Crecimiento Empresarial', bold=True)
add_para('Documento Interno — Versión 3.0 — Marzo 2026', italic=True)
add_para('Este documento define la estructura de precios, paquetes de servicio, definiciones de workflows, modalidades de contratación y estructura de compensación para Cresca.')

# --- Section 1 ---
doc.add_heading('1. POSICIONAMIENTO Y PROPUESTA DE VALOR', level=2)
add_para('Cresca NO es una agencia de marketing, una agencia de chatbots, ni un constructor de sitios web.')
add_para('Cresca ES una Empresa de Sistemas de Crecimiento Empresarial.', bold=True)
add_para('Ofrecemos infraestructura digital y automatización empresarial que permite a los negocios operar de manera más eficiente, capturar más clientes y escalar sus operaciones con sistemas — no con más empleados.')

doc.add_heading('Diferenciadores Clave', level=3)
add_para('• Sistemas completos, no servicios sueltos')
add_para('• Automatización que trabaja 24/7 sin costo adicional de personal')
add_para('• Enfoque en crecimiento estructurado y eficiencia operativa')
add_para('• Relación a largo plazo con soporte continuo')
add_para('• Todo integrado: web, CRM, automatización, chatbot, reservas')

# --- Section 2 ---
doc.add_heading('2. ESTRUCTURA DE PAQUETES — ESCALERA DE 3 NIVELES', level=2)
add_para('Cresca ofrece tres niveles de servicio diseñados para que los clientes comiencen pequeño y escalen cuando crezcan.', bold=True)

add_table(
    ['Nivel', 'Rango Mensual', 'Para Quién', 'Incluye'],
    [
        ['STARTER (Micro/Pequeño)', '$9.99 – $49.99', 'Negocios que inician su presencia digital', '1 Workflow Lite incluido'],
        ['GROWTH (Infraestructura)', '$50 – $100', 'Negocios listos para infraestructura + automatización', '1–2 Full Workflows incluidos'],
        ['PRO AUTOMATION (Avanzado)', 'Cotización', 'Operaciones avanzadas, workflows complejos', '4+ workflows, auditoría, monitoreo'],
    ]
)

# --- Section 2.1: STARTER ---
doc.add_heading('2.1 STARTER — Incluye 1 Workflow Lite', level=2)

doc.add_heading('Definición de Workflow Lite (no negociable)', level=3)
add_para('Un Workflow Lite es una automatización única, basada en plantilla, sin lógica compleja, seleccionada de un menú corto:')

doc.add_heading('Menú de Workflow Lite (elegir 1):', level=3)
add_para('1. Auto respuesta inicial (WhatsApp): "Gracias por escribir… horario… ubicación…"')
add_para('2. Mensaje de confirmación (cuando agendan/solicitan)')
add_para('3. Mensaje de recordatorio simple (trigger manual o programado)')
add_para('4. Solicitud de reseña (post-servicio) con link')
add_para('5. Mensaje de seguimiento 1 paso (24h después): "¿Desea agendar/completar pedido?"')

doc.add_heading('Reglas del Workflow Lite', level=3)
add_para('• Solo basado en plantilla (sin ramificación personalizada)')
add_para('• 1 canal solamente (WhatsApp O email)')
add_para('• 1 actualización/mes máximo')
add_para('• Sin integraciones más allá de herramientas básicas')
add_para('Esto mantiene Starter rentable y protege los márgenes.', italic=True)

doc.add_heading('Niveles Starter', level=3)

add_table(
    ['Tier', 'Precio', 'Qué Incluye', 'Setup Fee', 'Soporte'],
    [
        ['Starter $9.99\nDIY Lite', '$9.99/mes', 'Ebook + checklist (Google+FB)\nWhatsApp templates\nCopy templates\n1 Workflow Lite (solo plantilla)', '$0', 'Mínimo (solo mensajes)'],
        ['Starter $19.99\nStarter Assist', '$19.99/mes', 'Todo lo anterior +\n1 optimización Google Business (una vez)\n5 caption templates/mes\n1 Workflow Lite', '$0 – $49', 'Mínimo + email'],
        ['Starter $29.99\nStarter Presence', '$29.99/mes', 'Todo lo anterior +\nLink hub page O booking link setup (lite)\nCalendario de contenido básico\n1 Workflow Lite', '$49 – $99', 'Estándar'],
        ['Starter $49.99\nStarter Pro', '$49.99/mes', 'Todo lo anterior +\nMicro-sitio 1 página (plantilla)\n1 actualización mensual\n1 Workflow Lite', '$99 – $199', 'Estándar + 1 llamada/mes'],
    ]
)

# --- Section 2.2: GROWTH ---
doc.add_heading('2.2 GROWTH — Incluye Full Workflows', level=2)

doc.add_heading('Definición de Full Workflow', level=3)
add_para('Un Full Workflow incluye:')
add_para('• Trigger + acción + 1–2 pasos')
add_para('• Integración ligera (Sheets/CRM/email/WhatsApp)')
add_para('• Lógica básica permitida (if/then limitado)')

add_table(
    ['Tier', 'Precio', 'Qué Incluye', 'Setup Fee', 'Workflows'],
    [
        ['Growth $50\nFoundation', '$50/mes', 'Website template 3–5 páginas O módulo de reservas\nCaptura de leads → notificación\nGoogle + CTA buttons', '$199 – $299', '1 Workflow Lite'],
        ['Growth $75\nFoundation +1', '$75/mes', 'Todo lo anterior +\n1 Full Workflow module', '$299 – $399', '1 Full Workflow'],
        ['Growth $100\nFoundation +2', '$100/mes', 'Todo lo anterior +\n2 Full Workflow modules', '$399 – $499', '2 Full Workflows'],
    ]
)

# --- Section 2.3: PRO AUTOMATION ---
doc.add_heading('2.3 PRO AUTOMATION — Cotización', level=2)
add_para('Para negocios con operaciones complejas que necesitan automatización avanzada.', bold=True)
add_para('Incluye:')
add_para('• Auditoría completa + mapeo de procesos')
add_para('• 4+ workflows personalizados')
add_para('• Monitoreo y optimización continua')
add_para('• Workflows de operaciones internas, reportería, cobros, pipelines CRM')
add_para('• Recomendado compromiso de 12 meses')
add_para('Precio: Cotización personalizada basada en la complejidad del negocio.', bold=True)

# --- Section 3 ---
doc.add_heading('3. PRECIOS DE WORKFLOWS ADICIONALES (ADD-ON)', level=2)

doc.add_heading('Workflows Lite (Starter)', level=3)
add_para('• Incluido: 1 con cada plan Starter')
add_para('• Workflows Lite adicionales NO se venden por separado — el cliente debe actualizar a Growth.')

doc.add_heading('Full Workflows (Growth Add-Ons)', level=3)
add_para('Precio estándar: $15/mes por workflow adicional', bold=True)

add_table(
    ['Bundle', 'Precio Add-On', 'Ahorro'],
    [
        ['+2 workflows', '+$25/mes', '$5/mes de ahorro'],
        ['+5 workflows', '+$60/mes', '$15/mes de ahorro'],
        ['+8 workflows', '+$90/mes', '$30/mes de ahorro'],
    ]
)

# --- Section 4 ---
doc.add_heading('4. MÓDULO DE RESERVAS (CAPACIDAD UNIVERSAL)', level=2)
add_para('IMPORTANTE: Las reservas/booking NO son un add-on separado.', bold=True)
add_para('El módulo de reservas es una capacidad universal que puede incluirse dentro de cualquier paquete según la necesidad del cliente. Se implementa como parte del paquete Growth o Pro Automation cuando el negocio lo requiere.')
add_para('• Growth $50+: puede incluir módulo de booking como parte de la infraestructura')
add_para('• Pro Automation: integración completa con sistemas de reservas')

# --- Section 5 ---
doc.add_heading('5. ADD-ON: MARKETING DIGITAL', level=2)
add_para('Gestión operativa de campañas publicitarias en Meta Ads, Google Ads y otras plataformas.', bold=True)
add_para('• NO se garantiza ROI específico')
add_para('• Resultados dependen del mercado, competencia y presupuesto publicitario')
add_para('• Presupuesto de pauta publicitaria es responsabilidad del cliente y NO está incluido en la cuota')
add_para('• Solo se vende junto con un paquete principal (Growth o Pro Automation)')
add_para('Este es el único add-on disponible. El ad spend siempre es aparte.', italic=True)

# --- Section 6 ---
doc.add_heading('6. OPCIONES DE PAGO / PODER DE NEGOCIACIÓN', level=2)
add_para('Los representantes SOLO pueden ofrecer estas opciones estandarizadas:', bold=True)

add_table(
    ['Opción', 'Detalle', 'Incentivo'],
    [
        ['Mes a mes (estándar)', 'Sin compromiso de plazo', 'Precio estándar'],
        ['6 meses por adelantado', 'Pago de 6 meses upfront', '1 mes gratis O descuento en setup'],
        ['12 meses por adelantado', 'Pago de 12 meses upfront', '2 meses gratis O descuento en setup'],
        ['Plan Flex', 'Setup dividido: 50% hoy + 50% dentro de 30 días', 'Facilita el cierre sin descuento mensual'],
    ]
)
add_para('REGLA: No se permite descuento en la mensualidad a menos que sea pre-aprobado por la dirección.', bold=True)

# --- Section 7 ---
doc.add_heading('7. COMPENSACIÓN DEL REP FUNDADOR EN EL SALVADOR', level=2)
add_para('Este es el plan de compensación oficial para el primer representante en El Salvador:', bold=True)

add_table(
    ['Concepto', 'Detalle'],
    [
        ['Salario base', '$600/mes (garantizado)'],
        ['Comisión de setup', '20% de los fees de setup cobrados'],
        ['Comisión residual', '6% del ingreso recurrente mensual (MRR) cobrado'],
        ['Protección de calidad', 'Si un cliente cancela dentro de 60 días por mal fit/oversell → se recupera 50% de la comisión de setup de ese cliente (deducido del próximo pago)'],
    ]
)

doc.add_heading('Reglas de Pago', level=3)
add_para('• Las comisiones se pagan SOLO sobre dinero cobrado (no facturado)')
add_para('• El residual se paga SOLO mientras el cliente esté activo y al corriente')
add_para('• No hay comisión sobre costos de terceros (ad spend, fees de terceros, reembolsos)')
add_para('Ver documento separado: CRESCA_ES_Founding_Rep_Compensation.docx para detalles completos.', italic=True)

# --- Section 8 ---
doc.add_heading('8. ENFOQUE DE VENTAS', level=2)

doc.add_heading('Auditoría de Negocio en 15 Minutos', level=3)
add_para('El representante DEBE realizar una auditoría rápida de 15 minutos antes de recomendar un paquete:', bold=True)
add_para('1. Leads: ¿De dónde vienen sus clientes potenciales?')
add_para('2. Velocidad de respuesta: ¿Cuánto tardan en responder a un prospecto?')
add_para('3. Reservas/Ventas: ¿Cómo manejan el proceso de cierre?')
add_para('4. Operaciones: ¿Qué procesos internos consumen más tiempo?')
add_para('5. Retención: ¿Cómo mantienen contacto con clientes existentes?')

doc.add_heading('Árbol de Decisión', level=3)
add_table(
    ['Diagnóstico', 'Paquete Recomendado'],
    [
        ['No está listo digitalmente, necesita empezar', 'STARTER ($9.99–$49.99)'],
        ['Necesita reservas/leads + automatización simple', 'GROWTH ($50–$100)'],
        ['Dolor operacional + workflows complejos', 'PRO AUTOMATION (Cotización)'],
    ]
)

# --- Section 9 ---
doc.add_heading('9. ESTRATEGIA DE PRECIOS A FUTURO', level=2)

doc.add_heading('Cuándo Subir Precios', level=3)
add_para('• Cuando el producto mejora — cada nueva función de IA justifica un aumento del 10–15%')
add_para('• Cuando la demanda supere la capacidad')
add_para('• Incrementos máximos de 10% anual por contrato existente (con 30 días de aviso)')
add_para('• Anualmente — 5–8% es normal y raramente causa cancelaciones')

doc.add_heading('El Modelo "10x Rule"', level=3)
add_para('El valor que Cresca genera debe ser al menos 10x el costo del servicio. Si un negocio gasta $50/mes en Growth y el sistema le genera $500+ en ingresos adicionales o ahorros, el precio se justifica fácilmente.')

# --- Section 10 ---
doc.add_heading('10. RESUMEN EJECUTIVO', level=2)

add_table(
    ['Elemento', 'Detalle'],
    [
        ['Marca', 'Cresca — Sistemas de Crecimiento Empresarial'],
        ['Posicionamiento', 'Infraestructura + Automatización (no agencia de marketing)'],
        ['Paquete 1: STARTER', '$9.99–$49.99/mes — Productizado, bajo soporte, incluye 1 Workflow Lite'],
        ['Paquete 2: GROWTH', '$50–$100/mes — Infraestructura + automatización, 1–2 Full Workflows'],
        ['Paquete 3: PRO AUTOMATION', 'Cotización — Auditoría + 4+ workflows + monitoreo'],
        ['Add-On único', 'Marketing Digital (gestión; ad spend separado)'],
        ['Módulo de Reservas', 'Capacidad universal dentro de paquetes (no add-on separado)'],
        ['Workflow Lite adicional', 'No disponible — upgrade a Growth'],
        ['Full Workflow adicional', '$15/mes (o bundles: +2=$25, +5=$60, +8=$90)'],
        ['Pago', 'Mes a mes / 6 meses (1 free) / 12 meses (2 free) / Flex 50/50'],
        ['Rep ES Base', '$600/mes garantizado'],
        ['Comisión Setup', '20% de fees cobrados'],
        ['Comisión Residual', '6% del MRR cobrado'],
        ['Clawback', '50% setup si cancela en 60 días por mal fit'],
    ]
)

add_para('Cresca — Sistemas de Crecimiento Empresarial', bold=True)
add_para('Documento confidencial para uso interno. Versión 3.0 — Marzo 2026', italic=True)

doc.save(OUT)
print(f"Saved: {OUT}")
print("Done!")
