import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import pickle

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Colors
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0, 0, 0)
ORANGE = RGBColor(0xE0, 0x6C, 0x00)

# Configure styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    sname = f'Heading {level}'
    s = doc.styles[sname]
    s.font.name = 'Calibri'
    s.font.color.rgb = DARK_BLUE
    if level == 1:
        s.font.size = Pt(16)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after = Pt(6)
    elif level == 2:
        s.font.size = Pt(13)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(4)
    else:
        s.font.size = Pt(11)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(10)
        s.paragraph_format.space_after = Pt(3)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_body(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('☐  ' + item)
        run.font.size = Pt(9)

def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'{i}. {item}')
        run.font.size = Pt(9.5)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TODO DIRECTO')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('MANUAL DE PROCEDIMIENTOS ESTÁNDAR\nDE OPERACIONES Y ONBOARDING')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = MEDIUM_BLUE

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Guía Operativa Completa para Equipos Internos')
run.font.size = Pt(12)
run.italic = True

for _ in range(4):
    doc.add_paragraph()

# Version info table on title page
meta_data = [
    ['Versión', '1.0'],
    ['Fecha de Revisión', '28 de Febrero, 2026'],
    ['Clasificación', 'Confidencial – Solo Uso Interno'],
    ['Propietario', 'Todo Directo – Equipo de Operaciones'],
    ['Escalabilidad', 'Diseñado para 300+ clientes activos'],
]
t = doc.add_table(rows=len(meta_data), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for r, (label, val) in enumerate(meta_data):
    t.rows[r].cells[0].text = ''
    t.rows[r].cells[1].text = ''
    p0 = t.rows[r].cells[0].paragraphs[0]
    run0 = p0.add_run(label)
    run0.bold = True
    run0.font.size = Pt(10)
    p1 = t.rows[r].cells[1].paragraphs[0]
    run1 = p1.add_run(val)
    run1.font.size = Pt(10)
    set_cell_shading(t.rows[r].cells[0], 'D9E1F2')

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS PLACEHOLDER
# ============================================================
doc.add_heading('TABLA DE CONTENIDOS', level=1)
toc_items = [
    '1. Propósito y Alcance',
    '2. Roles y Responsabilidades',
    '3. Pipeline de Onboarding del Cliente (7 Etapas)',
    '   3.1 Cerrado / Pago Recibido',
    '   3.2 Materiales Pendientes',
    '   3.3 Construcción en Progreso',
    '   3.4 QA Interno',
    '   3.5 Revisión del Cliente',
    '   3.6 Activado',
    '   3.7 Seguimiento de 30 Días',
    '4. Checklists de Onboarding por Producto',
    '   4.1 AlToque App',
    '   4.2 Web + Booking (WaaS)',
    '   4.3 Vacation Rentals',
    '   4.4 Automatización de Procesos',
    '   4.5 Agente Chatbot con IA',
    '5. Reglas de Comunicación Interna',
    '6. Plantillas de Comunicación con el Cliente',
    '7. Reglas de Planificación de Capacidad',
    '8. Protocolo de Escalación',
    '9. Estrategia de Prevención de Churn',
    '10. Control de Calidad',
    '11. Documentación y Gestión de Archivos',
    '12. KPIs Operativos',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.size = Pt(10)

doc.add_page_break()

# Save intermediate
doc.save('.tmp/sop_ops_part1.docx')
print('✅ Part 1 complete: Title page + TOC')
