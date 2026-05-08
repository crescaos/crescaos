import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document('.tmp/sop_ops_part2.docx')
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')
    return table

def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('☐  ' + item)
        run.font.size = Pt(9)

def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'{i}. {item}')
        run.font.size = Pt(9.5)

# ============================================================
# SECTION 4: PRODUCT-SPECIFIC ONBOARDING CHECKLISTS
# ============================================================
doc.add_page_break()
doc.add_heading('4. CHECKLISTS DE ONBOARDING POR PRODUCTO', level=1)

p = doc.add_paragraph()
run = p.add_run('Cada producto tiene un checklist específico que el implementador debe completar antes de pasar a QA. Estos checklists complementan el checklist universal de QA de la Sección 3.4.')
run.font.size = Pt(10)

# 4.1 ALTOQUE
doc.add_heading('4.1 AlToque App', level=2)
p = doc.add_paragraph()
run = p.add_run('Plazo: 24–48 horas · Responsable: Implementador · Verifica: QA')
run.bold = True
run.font.size = Pt(9.5)
run.italic = True

add_checklist(doc, [
    'Catálogo de productos creado con fotos, nombres, descripciones y precios correctos',
    'Categorías organizadas lógicamente (mínimo 3 categorías para negocios con 20+ productos)',
    'Integración con WhatsApp Business verificada — pedido de prueba exitoso',
    'Flujo de pedido completo probado: selección → carrito → datos del cliente → envío por WhatsApp',
    'Logo y colores de marca del cliente aplicados',
    'Horarios de apertura/cierre configurados (si aplica)',
    'Enlace de la tienda funciona correctamente y es shareable',
    'Plan correcto aplicado (Chivo / Pro / Premium)',
    'Si es Pro/Premium: sin marca de agua, pagos integrados, reportes activos',
    'Prueba de pedido real enviada al WhatsApp del cliente como demo',
])

# 4.2 WEB + BOOKING
doc.add_heading('4.2 Web + Booking (WaaS)', level=2)
p = doc.add_paragraph()
run = p.add_run('Plazo: 5–14 días según plan · Responsable: Implementador · Verifica: QA')
run.bold = True
run.font.size = Pt(9.5)
run.italic = True

add_checklist(doc, [
    'Dominio configurado y apuntando correctamente (SSL/HTTPS activo)',
    'Página de inicio con información del negocio, fotos, y call-to-action claro',
    'Página de servicios con descripciones y precios (si autorizados)',
    'Formulario de contacto funcional — notificaciones llegan al email/WhatsApp del cliente',
    'Diseño responsive probado en iPhone, Android, tablet y desktop',
    'Google Analytics o equivalente instalado',
    'SEO básico: título de página, meta descripción, encabezados H1/H2, alt text en imágenes',
    'Velocidad de carga < 3 segundos (verificar con PageSpeed Insights)',
    'Si Plan Pro: sistema de reservas configurado con disponibilidad real del negocio',
    'Si Plan Pro: recordatorios automáticos (SMS/WhatsApp) programados y probados',
    'Si Plan Premium: e-commerce configurado con catálogo, carrito y pasarela de pagos',
    'Google Business Profile actualizado con link al nuevo sitio web',
    'Información de contacto correcta en todas las páginas (teléfono, dirección, horarios)',
    'Página de política de privacidad y términos (template estándar)',
])

# 4.3 VACATION RENTALS
doc.add_heading('4.3 Vacation Rentals', level=2)
p = doc.add_paragraph()
run = p.add_run('Plazo: 5–7 días · Responsable: Implementador · Verifica: QA')
run.bold = True
run.font.size = Pt(9.5)
run.italic = True

add_checklist(doc, [
    'Listado de propiedad completo: fotos (mín 15), descripción, amenidades, ubicación',
    'Calendario de disponibilidad sincronizado con otras plataformas (Airbnb, Booking) si aplica',
    'Precios configurados por temporada (alta, baja, feriados)',
    'Reglas de la casa publicadas y visibles',
    'Pasarela de pagos integrada y probada con transacción de prueba',
    'Flujo de reserva completo probado: búsqueda → selección → pago → confirmación',
    'Emails/mensajes de confirmación automática configurados',
    'Información de check-in/check-out clara y visible',
    'Galería de fotos optimizada para carga rápida',
    'Mapa de ubicación integrado',
    'Si Modelo B: sistema de gestión completa configurado (limpieza, llaves, comunicación)',
    'Comisión del modelo correcto configurada (8–10% Modelo A / 18–20% Modelo B)',
])

# 4.4 AUTOMATION
doc.add_heading('4.4 Automatización de Procesos', level=2)
p = doc.add_paragraph()
run = p.add_run('Plazo: 7–21 días según complejidad · Responsable: Implementador · Verifica: QA')
run.bold = True
run.font.size = Pt(9.5)
run.italic = True

add_checklist(doc, [
    'Documentación del flujo "antes" (proceso manual actual del cliente) archivada',
    'Flujos automatizados construidos y probados con datos reales/simulados',
    'Cada flujo tiene manejo de errores definido (qué pasa si falla un paso)',
    'Integraciones con sistemas del cliente configuradas (email, CRM, contabilidad, etc.)',
    'Notificaciones de éxito/fallo configuradas para el cliente y para el equipo interno',
    'Documentación de cada flujo: trigger, pasos, resultado esperado, tiempo de ejecución',
    'Prueba de extremo a extremo con datos reales del cliente (al menos 3 ejecuciones exitosas)',
    'Dashboard o reporte de monitoreo configurado para que el cliente vea resultados',
    'Plan de contingencia documentado si un flujo falla en producción',
    'Capacitación básica al cliente sobre cómo interpretar los reportes',
    'Si Pro: hasta 8 flujos verificados individualmente + en conjunto',
])

# 4.5 CHATBOT
doc.add_heading('4.5 Agente Chatbot con IA', level=2)
p = doc.add_paragraph()
run = p.add_run('Plazo: 5–14 días según plan · Responsable: Implementador · Verifica: QA')
run.bold = True
run.font.size = Pt(9.5)
run.italic = True

add_checklist(doc, [
    'Base de conocimiento del chatbot cargada (FAQ, servicios, precios, horarios)',
    'Tono de conversación configurado según la marca del cliente',
    'Respuestas probadas para al menos 20 preguntas frecuentes diferentes',
    'Flujo de escalación a humano configurado y probado (transferencia a WhatsApp/email del dueño)',
    'Canal principal integrado y activo (WhatsApp, sitio web, o ambos)',
    'Prueba de conversación completa simulando 5 escenarios distintos de usuario final',
    'Manejo de preguntas fuera de alcance probado (respuesta cortés + escalación)',
    'Idioma configurado (español por defecto, bilingüe español/inglés si requerido)',
    'Si Pro: múltiples canales activos y sincronizados',
    'Si Pro: funciones de pedido/reserva integradas y probadas',
    'Si Pro: reportes automáticos de conversaciones y métricas configurados',
    'Mensaje de bienvenida automático configurado',
    'Horarios de respuesta automática vs. respuesta con escalación definidos',
])

# ============================================================
# SECTION 5: INTERNAL COMMUNICATION RULES
# ============================================================
doc.add_page_break()
doc.add_heading('5. REGLAS DE COMUNICACIÓN INTERNA', level=1)

p = doc.add_paragraph()
run = p.add_run('La comunicación interna clara es la columna vertebral de las operaciones a escala. Estas reglas son obligatorias para todos los miembros del equipo.')
run.font.size = Pt(10)

doc.add_heading('5.1 Canales de Comunicación', level=2)

channels = [
    ['Canal Ops (Slack/WhatsApp)', 'Notificaciones de nuevos proyectos, cambios de estado en pipeline, alertas de SLA', 'Tiempo real — respuesta en < 1 hora durante horario laboral'],
    ['Canal Desarrollo', 'Discusiones técnicas, bloqueos, solicitudes de ayuda entre implementadores', 'Tiempo real durante construcción activa'],
    ['Canal Escalaciones', 'Problemas de cliente urgentes, fallos en producción, quejas', 'Respuesta obligatoria en < 30 minutos'],
    ['Email', 'Comunicaciones formales, contratos, documentación archivable', 'Respuesta en < 24 horas hábiles'],
    ['Standup Diario (10 min)', 'Revisión de pipeline, bloqueos del día, prioridades', 'Todos los días hábiles a las 8:30 AM'],
    ['Reunión Semanal (30 min)', 'KPIs de la semana, retrospectiva, planificación', 'Todos los lunes a las 9:00 AM'],
]

make_table(doc, ['Canal', 'Uso', 'SLA de Respuesta'], channels)

doc.add_heading('5.2 Reglas de Comunicación', level=2)

add_numbered(doc, [
    'Formato estándar para reportar nuevo proyecto: "NUEVO: [Producto] — [Cliente] — [Rep] — [Fecha de cierre]".',
    'Formato para cambio de estado: "ESTADO: [Cliente] — [Etapa anterior] → [Etapa nueva] — [Notas]".',
    'Formato para escalación: "⚠️ ESCALAR: [Cliente] — [Tipo de problema] — [Urgencia: Alta/Media/Baja]".',
    'Nunca comunicar información de cliente en canales públicos o no autorizados.',
    'Toda decisión importante debe documentarse por escrito, no solo verbal.',
    'Si un proyecto está bloqueado, el responsable debe comunicarlo en < 2 horas, no esperar al standup.',
    'Los cambios de alcance solicitados por el cliente se documentan y aprueban antes de implementar.',
    'No se aceptan nuevos proyectos sin formulario de intake completo — sin excepciones.',
])

# ============================================================
# SECTION 6: CLIENT COMMUNICATION TEMPLATES
# ============================================================
doc.add_heading('6. PLANTILLAS DE COMUNICACIÓN CON EL CLIENTE', level=1)

p = doc.add_paragraph()
run = p.add_run('Todas las comunicaciones con el cliente deben seguir estas plantillas para mantener consistencia y profesionalismo. Personalizar con el nombre del cliente y detalles específicos.')
run.font.size = Pt(10)

templates = [
    ['Bienvenida Post-Cierre',
     '¡Hola [Nombre]! 🎉 Bienvenido/a a Todo Directo. Su proyecto de [Producto] ya está en proceso. Nuestro equipo comenzará a trabajar en cuanto recibamos los materiales necesarios. Le enviaré la lista en un momento. Cualquier duda, estoy a sus órdenes. — [Rep]'],
    ['Solicitud de Materiales',
     '¡Hola [Nombre]! Para iniciar la construcción de su [Producto] necesitamos lo siguiente:\n[Lista de materiales]\nPuede enviárnoslo por este medio o a [email]. Mientras más rápido los recibamos, más rápido estará listo. ¿Alguna pregunta?'],
    ['Recordatorio de Materiales (48h)',
     '¡Hola [Nombre]! Solo un recordatorio amable — aún necesitamos [materiales pendientes] para avanzar con su proyecto. ¿Le puedo ayudar con algo para facilitar la entrega?'],
    ['Producto Listo para Revisión',
     '¡Hola [Nombre]! 🚀 Tenemos listo su [Producto] para que lo revise. [Link de acceso]. Tómese un momento para revisarlo y me dice si hay algo que le gustaría ajustar. Tenemos hasta 2 rondas de cambios incluidas.'],
    ['Activación Exitosa',
     '¡Hola [Nombre]! 🎉 Su [Producto] ya está ACTIVO y funcionando.\n[Link de acceso / instrucciones]\nRecuerde que estoy disponible para cualquier duda. Le contactaré en una semana para ver cómo le va. ¡Éxito!'],
    ['Seguimiento Día 7',
     '¡Hola [Nombre]! Ya lleva una semana con su [Producto]. ¿Cómo le ha ido? ¿Ha recibido [pedidos/reservas/consultas]? ¿Alguna duda que pueda resolver?'],
    ['Seguimiento Día 30',
     '¡Hola [Nombre]! Ya cumplimos un mes juntos. 📊 Aquí está su reporte del primer mes:\n• [Métrica 1]\n• [Métrica 2]\n• [Métrica 3]\n¿Le gustaría explorar cómo podemos mejorar estos números con [producto complementario]?'],
    ['Prevención de Cancelación',
     '¡Hola [Nombre]! Entendemos que tiene algunas inquietudes sobre su [Producto]. Me gustaría agendar una llamada breve para entender mejor su situación y ver cómo podemos ayudarle. ¿Tiene disponibilidad [fecha/hora]?'],
]

for name, template in templates:
    doc.add_heading(f'Plantilla: {name}', level=3)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(template)
    run.font.size = Pt(9)
    run.italic = True

doc.save('.tmp/sop_ops_part3.docx')
print('✅ Part 3 complete: Sections 4-6 (Product Checklists, Communication)')
