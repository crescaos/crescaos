"""Generate Operations SOP Part 2 (Sections 7-12) and merge with Part 1 — 3-Tier Model."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

# Load part 1
doc = Document(r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\.tmp\sop_part1.docx')

def ap(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def at(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.text = ''
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i+1, j)
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    doc.add_paragraph()
    return t

# === SECTION 7 ===
doc.add_heading('SECCIÓN 7: REGLAS DE PLANIFICACIÓN DE CAPACIDAD', level=1)

doc.add_heading('Regla Fundamental: Capacidad Máxima', level=2)
ap('El fundador puede gestionar un número limitado de proyectos activos simultáneamente. Esta regla es innegociable hasta que se contrate equipo técnico adicional.', bold=True)

doc.add_heading('Límites de Capacidad por Paquete', level=2)
at(['Paquete', 'Complejidad', 'Tiempo Promedio', 'Capacidad Simultánea'], [
    ['STARTER ($9.99–$49.99)', 'Baja', '1–5 días', 'Alto volumen posible'],
    ['GROWTH ($50–$100)', 'Media', '1–3 semanas', '3–5 simultáneos'],
    ['PRO AUTOMATION', 'Alta', '4–8 semanas', '1–2 simultáneos'],
])

doc.add_heading('Reglas de Expansión de Capacidad', level=2)
ap('Contratar primer técnico cuando:', bold=True)
ap('• 80%+ de capacidad por 2+ semanas consecutivas')
ap('• MRR supere $5,000/mes')
ap('• Cola de espera de nuevos proyectos > 2 semanas')

# === SECTION 8 ===
doc.add_heading('SECCIÓN 8: PROTOCOLO DE ESCALAMIENTO', level=1)

doc.add_heading('Niveles de Escalamiento', level=2)
at(['Nivel', 'Descripción', 'Quién Resuelve', 'Tiempo Máximo'], [
    ['Nivel 1', 'Problemas técnicos menores, preguntas de uso', 'Soporte', '4 horas'],
    ['Nivel 2', 'Problemas técnicos complejos, integraciones', 'Técnico Senior', '24 horas'],
    ['Nivel 3', 'Fallos críticos, cliente en riesgo de cancelar', 'Fundador', '2 horas'],
])

doc.add_heading('Cuándo Escalar Automáticamente', level=2)
ap('• Cliente no responde hace más de 10 días hábiles → Nivel 2')
ap('• Sistema del cliente caído por más de 4 horas → Nivel 3')
ap('• Cliente menciona cancelar o está insatisfecho → Nivel 3')
ap('• Problema técnico no resuelto en 48 horas → Nivel 2')

doc.add_heading('Formato de Reporte de Escalamiento', level=2)
ap('📣 ESCALAMIENTO — Nivel [X]', bold=True)
ap('Cliente: [Nombre]')
ap('Paquete: [STARTER / GROWTH / PRO AUTOMATION]')
ap('Problema: [Descripción breve]')
ap('Impacto: [Qué está afectado]')
ap('Acciones tomadas: [Qué se intentó]')
ap('Acción requerida: [Qué necesitamos]')

# === SECTION 9 ===
doc.add_heading('SECCIÓN 9: ESTRATEGIA DE PREVENCIÓN DE CHURN', level=1)

doc.add_heading('Señales de Riesgo de Cancelación', level=2)
at(['Señal', 'Riesgo', 'Acción Inmediata'], [
    ['No usa los sistemas hace 2+ semanas', 'Alto', 'Llamar para ofrecer re-capacitación'],
    ['Quejas repetidas sobre el mismo tema', 'Alto', 'Escalar a Nivel 3, reunión con Fundador'],
    ['Pagos retrasados de forma recurrente', 'Medio', 'Conversación sobre expectativas y valor'],
    ['Pregunta por reducir servicios', 'Medio', 'Ofrecer bajar de tier antes de cancelar'],
    ['No responde a comunicaciones', 'Medio', 'Contactar por canal alternativo'],
])

doc.add_heading('Protocolo de Retención', level=2)
ap('Cuando un cliente manifiesta querer cancelar:', bold=True)
ap('1. NO ofrezca descuento inmediatamente.')
ap('2. Pregunte la razón real.')
ap('3. Si es resoluble, proponga plan de acción.')
ap('4. Si es presupuestario, opciones autorizadas:')
ap('   • Bajar a tier inferior (ej: Growth → Starter Pro)')
ap('   • Ofrecer Plan Flex si es problema de flujo de caja')
ap('   • Pausa de 1 mes sin penalidad (máximo una vez)')
ap('5. Si cancela, hacer entrevista de salida y documentar en CRM.')

# === SECTION 10 ===
doc.add_heading('SECCIÓN 10: CHECKLIST DE CONTROL DE CALIDAD', level=1)

doc.add_heading('QA Universal', level=2)
ap('☐ Todos los links funcionan correctamente')
ap('☐ Diseño responsivo verificado (móvil + desktop)')
ap('☐ Sin errores ortográficos en contenido visible')
ap('☐ Workflow probado end-to-end')
ap('☐ Notificaciones activadas y probadas')

doc.add_heading('QA por Paquete', level=2)
doc.add_heading('STARTER', level=3)
ap('☐ Workflow Lite disparándose correctamente')
ap('☐ Plantillas entregadas y verificadas')
ap('☐ Google Business optimizado (si aplica)')

doc.add_heading('GROWTH', level=3)
ap('☐ Sitio web funcional en móvil y desktop')
ap('☐ SEO básico configurado')
ap('☐ Captura de leads probada')
ap('☐ Full Workflow(s) probado(s) end-to-end')
ap('☐ Módulo de reservas probado (si aplica)')
ap('☐ Google Analytics verificado')

doc.add_heading('PRO AUTOMATION', level=3)
ap('☐ Cada automatización probada end-to-end')
ap('☐ Integraciones transmitiendo datos correctamente')
ap('☐ Reportes generando datos precisos')
ap('☐ Documentación de procesos entregada')

# === SECTION 11 ===
doc.add_heading('SECCIÓN 11: ESTÁNDARES DE DOCUMENTACIÓN', level=1)
ap('Formato de nombres: [NombreCliente]_[TipoDocumento]_[Fecha].ext', bold=True)
ap('Ejemplo: CafeAroma_Contrato_2026-03-01.pdf')

at(['Documento', 'Cuándo se Crea', 'Responsable'], [
    ['Contrato firmado', 'Al cierre', 'Rep Ventas'],
    ['Formulario de intake completado', 'Pre-cierre o día 1', 'Rep Ventas'],
    ['Acta de kickoff', 'Reunión de kickoff', 'Técnico'],
    ['Checklist de QA completado', 'Pre go-live', 'Técnico'],
    ['Registro de pago mensual', 'Cada mes', 'Finanzas'],
])

# === SECTION 12 ===
doc.add_heading('SECCIÓN 12: KPIs OPERACIONALES', level=1)

at(['KPI', 'Meta', 'Frecuencia'], [
    ['Implementación a tiempo', '≥ 85%', 'Mensual'],
    ['Satisfacción (NPS)', '≥ 8/10', 'Trimestral'],
    ['Tiempo de onboarding', 'STARTER ≤5d, GROWTH ≤3sem, PRO ≤8sem', 'Mensual'],
    ['Tasa de churn', '≤ 5%', 'Mensual'],
    ['Tickets resueltos en SLA', '≥ 90%', 'Mensual'],
    ['MRR', 'Creciente', 'Mensual'],
])

ap('')
ap('Cresca — Sistemas de Crecimiento Empresarial', bold=True)
ap('Manual de Operaciones — Documento confidencial. Versión 3.0 — Marzo 2026', italic=True)

# Save final document
OUT = r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\CRESCA_Operations_SOP.docx'
doc.save(OUT)
print(f"Saved: {OUT}")
print("Done!")
