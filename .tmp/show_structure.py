import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

BASE = r'c:\Users\12132\Desktop\DigiFacil\DigiFacil'

for fname in ['Todo_Directo_Pricing_Strategy.docx','Todo_Directo_Operations_SOP.docx','Todo_Directo_SOP_v2.docx','Todo_Directo_2_Paginas_Sales_Cheat_Sheet_ES_v2.docx','Todo_Directo_Intake_Form.docx']:
    doc = Document(os.path.join(BASE, fname))
    headings = []
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith('Heading') and p.text.strip():
            headings.append(f"  [{p.style.name}] {p.text.strip()}")
    num_para = len([p for p in doc.paragraphs if p.text.strip()])
    num_tables = len(doc.tables)
    print(f"\n=== {fname} === ({num_para} paragraphs, {num_tables} tables)")
    for h in headings:
        print(h)
