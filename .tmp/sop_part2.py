import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document('.tmp/sop_ops_part1.docx')

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')
    return table

def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('☐  ' + item)
        run.font.size = Pt(9)

def add_numbered(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f'{i}. {item}')
        run.font.size = Pt(9.5)

# ============================================================
# SECTION 1: PURPOSE AND SCOPE
# ============================================================
doc.add_heading('1. PROPÓSITO Y ALCANCE', level=1)

p = doc.add_paragraph()
run = p.add_run('Este Manual de Procedimientos Estándar de Operaciones (SOP) define los procesos, flujos de trabajo, acuerdos de nivel de servicio (SLAs), controles de calidad y métricas operativas que rigen el onboarding y la gestión de clientes de Todo Directo.')
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Objetivos principales:')
run.bold = True
run.font.size = Pt(10)

add_numbered(doc, [
    'Estandarizar el proceso de onboarding para los cinco productos de Todo Directo.',
    'Garantizar tiempos de entrega predecibles y cumplibles para cada producto.',
    'Escalar operaciones de manera sostenible hasta 300+ clientes activos.',
    'Minimizar el churn mediante seguimiento proactivo y control de calidad.',
    'Crear trazabilidad completa en cada etapa del pipeline.',
    'Establecer métricas claras para medir la salud operativa del negocio.',
])

p = doc.add_paragraph()
run = p.add_run('Alcance: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('Este documento aplica a todos los miembros del equipo interno de Todo Directo, incluyendo representantes comerciales, equipo de implementación, soporte técnico y liderazgo. No reemplaza el SOP de Ventas (documento separado) pero se complementa con él.')
run.font.size = Pt(10)

# ============================================================
# SECTION 2: ROLES AND RESPONSIBILITIES
# ============================================================
doc.add_heading('2. ROLES Y RESPONSABILIDADES', level=1)

p = doc.add_paragraph()
run = p.add_run('Cada rol tiene responsabilidades específicas dentro del pipeline de onboarding. La claridad en roles es crítica para evitar cuellos de botella a escala.')
run.font.size = Pt(10)

roles = [
    ['Representante Comercial (Rep)', 'Cierre de venta, recolección de materiales del cliente, llenado del formulario de intake, entrega al equipo de implementación. Primer punto de contacto del cliente.'],
    ['Líder de Implementación', 'Asignación de proyectos al equipo técnico, seguimiento de SLAs, comunicación de actualizaciones al rep y al cliente. Responsable del pipeline completo.'],
    ['Desarrollador / Implementador', 'Construcción técnica del producto (sitio web, app, chatbot, automatización). Cumplimiento de los estándares de calidad definidos en este documento.'],
    ['Especialista QA', 'Revisión de calidad antes de la entrega al cliente. Verificación de funcionalidad, diseño, rendimiento y cumplimiento de especificaciones. Puede ser el Líder de Implementación en equipos pequeños.'],
    ['Soporte al Cliente', 'Gestión de tickets post-activación, soporte técnico nivel 1, escalación de problemas técnicos al equipo de implementación.'],
    ['Director de Operaciones', 'Supervisión del pipeline completo, revisión de KPIs semanales, resolución de escalaciones nivel 3, planificación de capacidad y contratación.'],
]

make_table(doc, ['Rol', 'Responsabilidades'], roles)

doc.add_heading('2.1 Matriz RACI por Etapa', level=2)

p = doc.add_paragraph()
run = p.add_run('R = Responsable (ejecuta) · A = Aprobador · C = Consultado · I = Informado')
run.italic = True
run.font.size = Pt(9)

raci = [
    ['Cierre / Pago', 'R', 'I', '–', '–', '–', 'I'],
    ['Materiales Pendientes', 'R', 'A', '–', '–', '–', 'I'],
    ['Construcción', 'I', 'A', 'R', '–', '–', 'I'],
    ['QA Interno', 'I', 'C', 'C', 'R', '–', 'I'],
    ['Revisión Cliente', 'R', 'A', 'C', '–', '–', 'I'],
    ['Activación', 'I', 'R', 'C', '–', 'I', 'I'],
    ['Seguimiento 30 Días', 'R', 'I', '–', '–', 'C', 'I'],
]

make_table(doc, ['Etapa', 'Rep', 'Líder Impl.', 'Dev', 'QA', 'Soporte', 'Director'], raci)

# ============================================================
# SECTION 3: PIPELINE DE ONBOARDING (7 STAGES)
# ============================================================
doc.add_heading('3. PIPELINE DE ONBOARDING DEL CLIENTE', level=1)

p = doc.add_paragraph()
run = p.add_run('El pipeline de onboarding consta de 7 etapas secuenciales. Cada etapa tiene un SLA definido, un responsable, y criterios de entrada/salida claros. Ningún proyecto avanza a la siguiente etapa sin cumplir los criterios de salida.')
run.font.size = Pt(10)

# Stage overview table
stages_overview = [
    ['1. Cerrado / Pago Recibido', 'Rep Comercial', '24 horas', 'Pago confirmado + formulario de intake completo'],
    ['2. Materiales Pendientes', 'Rep + Cliente', '72 horas máx', 'Todos los materiales requeridos entregados'],
    ['3. Construcción en Progreso', 'Implementador', 'Ver tabla por producto', 'Producto construido y funcionando internamente'],
    ['4. QA Interno', 'Especialista QA', '24–48 horas', 'Checklist de QA aprobado sin errores'],
    ['5. Revisión del Cliente', 'Rep + Cliente', '48 horas', 'Cliente aprueba o solicita cambios (máx 2 rondas)'],
    ['6. Activado', 'Líder Impl.', '24 horas post-aprobación', 'Producto en producción + cliente notificado'],
    ['7. Seguimiento 30 Días', 'Rep Comercial', 'Día 7, 14, 30', 'Satisfacción confirmada + oportunidades de upsell'],
]

make_table(doc, ['Etapa', 'Responsable', 'SLA', 'Criterio de Salida'], stages_overview)

# --- STAGE 3.1 ---
doc.add_heading('3.1 Etapa 1: Cerrado / Pago Recibido', level=2)

p = doc.add_paragraph()
run = p.add_run('SLA: 24 horas desde el cierre para entregar al equipo de implementación.')
run.bold = True
run.font.size = Pt(9.5)

add_numbered(doc, [
    'El representante confirma el pago recibido (efectivo, transferencia, Chivo Wallet o tarjeta).',
    'El representante completa el Formulario de Orden de Servicio (Intake Form) con toda la información del cliente.',
    'El representante envía el formulario completo al canal interno designado (Slack/WhatsApp Ops) dentro de las 24 horas.',
    'El Líder de Implementación confirma recepción y asigna el proyecto en el tracker.',
    'El cliente recibe un mensaje de bienvenida automatizado confirmando que su proyecto está en proceso.',
])

p = doc.add_paragraph()
run = p.add_run('Criterio de salida: ')
run.bold = True
run.font.size = Pt(9.5)
run = p.add_run('Pago registrado + Formulario completo + Proyecto asignado en tracker + Mensaje de bienvenida enviado.')
run.font.size = Pt(9.5)

# --- STAGE 3.2 ---
doc.add_heading('3.2 Etapa 2: Materiales Pendientes', level=2)

p = doc.add_paragraph()
run = p.add_run('SLA: Máximo 72 horas. Si los materiales no se reciben en 72 horas, el rep escala con llamada telefónica.')
run.bold = True
run.font.size = Pt(9.5)

p = doc.add_paragraph()
run = p.add_run('Materiales requeridos por producto:')
run.bold = True
run.font.size = Pt(10)

materials = [
    ['AlToque App', 'Logo, fotos de productos (mín. 10), lista de precios, número de WhatsApp, descripción del negocio'],
    ['Web + Booking', 'Logo, fotos del negocio (mín. 5), descripción de servicios, horarios, información de contacto, datos para Google Business'],
    ['Vacation Rentals', 'Fotos de la propiedad (mín. 15), precios por noche/temporada, calendario de disponibilidad, reglas de la casa, ubicación exacta'],
    ['Automatización', 'Documentación de flujos actuales, acceso a sistemas existentes, lista de integraciones requeridas, ejemplos de reportes actuales'],
    ['Chatbot IA', 'Lista de preguntas frecuentes (mín. 20), tono de marca, horario de atención, información de servicios/productos, flujo de escalación a humano'],
]

make_table(doc, ['Producto', 'Materiales Requeridos'], materials)

add_numbered(doc, [
    'El rep envía al cliente la lista de materiales y explica cómo entregarlos (WhatsApp, email, Google Drive).',
    'Primer recordatorio a las 24 horas si no se han recibido.',
    'Segundo recordatorio a las 48 horas con llamada telefónica.',
    'Si a las 72 horas no se reciben, el Líder de Implementación contacta al rep para evaluar si el cliente está comprometido.',
    'Máximo 7 días en esta etapa antes de marcar como "Estancado" y aplicar protocolo de rescate.',
])

# --- STAGE 3.3 ---
doc.add_heading('3.3 Etapa 3: Construcción en Progreso', level=2)

p = doc.add_paragraph()
run = p.add_run('SLAs de construcción por producto:')
run.bold = True
run.font.size = Pt(10)

build_slas = [
    ['AlToque App', '24–48 horas', 'Catálogo cargado, pedidos funcionales, WhatsApp integrado'],
    ['Web Plan Chivo', '5–7 días hábiles', 'Página de inicio, formulario de contacto, móvil-friendly'],
    ['Web Plan Pro', '7–10 días hábiles', 'Todo Chivo + reservas + recordatorios SMS'],
    ['Web Plan Premium', '10–14 días hábiles', 'Todo Pro + e-commerce + soporte VIP'],
    ['Vacation Rentals', '5–7 días hábiles', 'Listado completo, calendario de reservas, pasarela de pagos'],
    ['Automatización Básica', '7–14 días hábiles', '1–3 flujos automatizados funcionando'],
    ['Automatización Pro', '14–21 días hábiles', 'Hasta 8 flujos + integraciones + reportes'],
    ['Chatbot Básico', '5–7 días hábiles', '1 canal activo + FAQ programadas'],
    ['Chatbot Pro', '7–14 días hábiles', 'Múltiples canales + pedidos/reservas + reportes'],
]

make_table(doc, ['Producto', 'Plazo SLA', 'Entregable'], build_slas)

add_numbered(doc, [
    'El implementador inicia la construcción dentro de las 24 horas de recibir materiales completos.',
    'El Líder de Implementación actualiza el tracker con el estado de progreso cada 48 horas.',
    'Si el plazo SLA se va a exceder, notificar al rep y al cliente con 48 horas de anticipación y proporcionar nueva fecha estimada.',
    'Los proyectos en construcción se revisan en el standup diario del equipo técnico.',
    'El implementador notifica al Líder de Implementación cuando el producto está listo para QA.',
])

# --- STAGE 3.4 ---
doc.add_heading('3.4 Etapa 4: QA Interno', level=2)

p = doc.add_paragraph()
run = p.add_run('SLA: 24–48 horas. Ningún producto se entrega al cliente sin pasar QA interno.')
run.bold = True
run.font.size = Pt(9.5)

p = doc.add_paragraph()
run = p.add_run('Checklist Universal de QA (aplicable a todos los productos):')
run.bold = True
run.font.size = Pt(10)

add_checklist(doc, [
    'Funcionalidad principal opera correctamente (pedidos, reservas, flujos, respuestas)',
    'Diseño visual coincide con los estándares de marca del cliente',
    'Totalmente responsive (móvil, tablet, desktop)',
    'Velocidad de carga aceptable (< 3 segundos)',
    'Sin errores de ortografía ni contenido placeholder',
    'Links funcionales — ningún enlace roto',
    'Formularios de contacto probados y funcionando',
    'Integración de WhatsApp verificada',
    'SEO básico configurado (títulos, meta descripciones)',
    'SSL activo (HTTPS)',
    'Información del cliente correcta (nombre, dirección, teléfono)',
    'Prueba de flujo completo de usuario final (pedido/reserva/consulta)',
])

add_numbered(doc, [
    'El QA revisa el producto contra el checklist completo (universal + específico del producto).',
    'Si encuentra errores, los documenta y devuelve al implementador con descripción clara.',
    'El implementador corrige en un máximo de 24 horas.',
    'QA re-verifica solo los puntos fallidos.',
    'Una vez aprobado, el QA marca el proyecto como "Listo para Cliente" y notifica al rep.',
])

# --- STAGE 3.5 ---
doc.add_heading('3.5 Etapa 5: Revisión del Cliente', level=2)

p = doc.add_paragraph()
run = p.add_run('SLA: El cliente tiene 48 horas para revisar y responder. Máximo 2 rondas de cambios.')
run.bold = True
run.font.size = Pt(9.5)

add_numbered(doc, [
    'El rep presenta el producto al cliente (en persona, por llamada o por video).',
    'El rep captura el feedback del cliente de forma estructurada: qué cambiar, qué agregar, qué quitar.',
    'Cambios menores (texto, fotos, colores) se implementan en 24 horas.',
    'Cambios de alcance (nuevas funciones no cotizadas) requieren cotización adicional.',
    'Si el cliente no responde en 48 horas, el rep hace seguimiento por llamada.',
    'Después de 2 rondas de cambios, cualquier modificación adicional se cobra como soporte.',
    'El cliente da aprobación verbal o escrita (WhatsApp/email). Se documenta en el tracker.',
])

# --- STAGE 3.6 ---
doc.add_heading('3.6 Etapa 6: Activado', level=2)

p = doc.add_paragraph()
run = p.add_run('SLA: 24 horas después de la aprobación del cliente, el producto está en producción.')
run.bold = True
run.font.size = Pt(9.5)

add_numbered(doc, [
    'El Líder de Implementación ejecuta el lanzamiento a producción (deploy, activación de dominio, publicación).',
    'Se configura el monitoreo automático del producto (uptime, errores).',
    'El rep envía al cliente el mensaje de "Activación Exitosa" con instrucciones de uso.',
    'Se programa automáticamente el seguimiento a los 7, 14 y 30 días.',
    'El proyecto se marca como "Activo" en el tracker y en el CRM/Cotizador.',
    'El ingreso recurrente se activa en el sistema de facturación.',
])

# --- STAGE 3.7 ---
doc.add_heading('3.7 Etapa 7: Seguimiento de 30 Días', level=2)

p = doc.add_paragraph()
run = p.add_run('SLA: Contacto obligatorio a los 7, 14 y 30 días post-activación.')
run.bold = True
run.font.size = Pt(9.5)

followup = [
    ['Día 7', 'Rep Comercial', 'WhatsApp o llamada', '"¿Cómo le ha ido con [producto]? ¿Todo funcionando bien?"', 'Resolver cualquier duda inicial, confirmar que el cliente está usando el producto.'],
    ['Día 14', 'Rep Comercial', 'Llamada', 'Revisión de métricas iniciales (pedidos, reservas, consultas). Identificar si hay problemas.', 'Si no está usando el producto, activar protocolo de rescate.'],
    ['Día 30', 'Rep + Soporte', 'Llamada + Reporte', 'Entrega de reporte de primer mes. Explorar oportunidades de upsell/stack.', 'Este es el momento clave de retención. Si el cliente ve valor, se queda.'],
]

make_table(doc, ['Hito', 'Responsable', 'Canal', 'Acción', 'Objetivo'], followup)

doc.save('.tmp/sop_ops_part2.docx')
print('✅ Part 2 complete: Sections 1-3 (Purpose, Roles, Pipeline)')
