"""Generate the updated Cresca Operations SOP - Part 1 (Sections 1-6) — 3-Tier Model."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

sn = doc.styles['Normal']
sn.font.name = 'Calibri'
sn.font.size = Pt(11)
sn.paragraph_format.space_after = Pt(4)

for lvl in [1,2,3]:
    st = doc.styles[f'Heading {lvl}']
    st.font.name = 'Calibri'
    st.font.bold = True
    st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    st.font.size = Pt([0,18,14,12][lvl])

def ap(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def at(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.text = ''
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.cell(i+1, j)
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    doc.add_paragraph()
    return t

# ======================= CONTENT =======================

doc.add_heading('CRESCA — Manual de Operaciones (SOP)', level=1)
ap('Sistemas de Crecimiento Empresarial', bold=True)
ap('Documento Interno — Versión 3.0 — Marzo 2026', italic=True)

# --- TOC ---
doc.add_heading('TABLA DE CONTENIDOS', level=1)
sections = [
    'Sección 1: Propósito y Alcance',
    'Sección 2: Roles y Responsabilidades',
    'Sección 3: Pipeline de Onboarding de Clientes',
    'Sección 4: Checklists de Onboarding por Paquete',
    'Sección 5: Reglas de Comunicación Interna',
    'Sección 6: Plantillas de Comunicación con el Cliente',
    'Sección 7: Reglas de Planificación de Capacidad',
    'Sección 8: Protocolo de Escalamiento',
    'Sección 9: Estrategia de Prevención de Churn',
    'Sección 10: Checklist de Control de Calidad',
    'Sección 11: Estándares de Documentación y Archivos',
    'Sección 12: KPIs Operacionales',
]
for s in sections:
    ap(s)

# === SECTION 1 ===
doc.add_heading('SECCIÓN 1: PROPÓSITO Y ALCANCE', level=1)
ap('Este manual establece los procedimientos operativos estándar para todas las operaciones de Cresca, incluyendo onboarding de clientes, implementación de sistemas, control de calidad, comunicación y gestión de capacidad.')

doc.add_heading('¿Quién Debe Usar Este Manual?', level=2)
ap('• Fundador / Director de Operaciones')
ap('• Equipo técnico de implementación')
ap('• Representantes de ventas (secciones relevantes)')
ap('• Cualquier colaborador involucrado en la entrega de servicios')

doc.add_heading('¿Qué NO Cubre Este Manual?', level=2)
ap('• Estrategia de precios (ver: Cresca Pricing Strategy)')
ap('• Scripts de ventas (ver: Cresca Sales SOP)')
ap('• Aspectos legales/contractuales (ver: Contrato de Servicios)')
ap('• Compensación del rep fundador (ver: CRESCA_ES_Founding_Rep_Compensation.docx)')

doc.add_heading('Control de Versiones', level=2)
at(['Versión', 'Fecha', 'Cambios'], [
    ['1.0', 'Enero 2026', 'Versión inicial (Todo Directo)'],
    ['2.0', 'Marzo 2026', 'Rebrand a Cresca. 2 paquetes + 2 add-ons. Comisiones 25%/8%.'],
    ['3.0', 'Marzo 2026', 'Nuevo modelo 3-tier (STARTER/GROWTH/PRO AUTOMATION). Comisiones 20%/6%. Booking como módulo universal. Rep fundador $600 base.'],
])

# === SECTION 2 ===
doc.add_heading('SECCIÓN 2: ROLES Y RESPONSABILIDADES', level=1)

doc.add_heading('Organigrama Funcional', level=2)
at(['Rol', 'Responsabilidades Principales'], [
    ['Fundador / CEO', 'Estrategia, cierre de ventas Pro Automation, supervisión general, builds activos'],
    ['Rep Fundador (ES)', 'Prospección, auditoría de 15 min, calificación, cierre de ventas Starter/Growth, seguimiento'],
    ['Técnico de Implementación', 'Configuración de sistemas, workflows, integraciones, QA'],
    ['Soporte al Cliente', 'Atención post-lanzamiento, tickets, seguimiento de satisfacción'],
])

doc.add_heading('Matriz RACI — Procesos Clave', level=2)
at(['Proceso', 'Fundador', 'Rep Ventas', 'Técnico', 'Soporte'], [
    ['Captación de leads', 'I', 'R', '—', '—'],
    ['Auditoría 15 min', 'C', 'R', '—', '—'],
    ['Calificación', 'C', 'R', '—', '—'],
    ['Cierre de venta', 'A', 'R', 'I', '—'],
    ['Onboarding', 'A', 'I', 'R', 'I'],
    ['Implementación', 'A', '—', 'R', '—'],
    ['QA y pruebas', 'A', '—', 'R', 'C'],
    ['Go-Live', 'A', 'I', 'R', 'R'],
    ['Soporte continuo', 'I', '—', 'C', 'R'],
    ['Retención / Churn', 'A', 'R', 'C', 'R'],
])
ap('R=Responsable, A=Aprueba, C=Consultado, I=Informado', italic=True)

# === SECTION 3 ===
doc.add_heading('SECCIÓN 3: PIPELINE DE ONBOARDING DE CLIENTES', level=1)

doc.add_heading('Vista General del Pipeline', level=2)
at(['Etapa', 'Descripción', 'Responsable', 'Duración Típica'], [
    ['1. Captura del Lead', 'Registro del prospecto en CRM', 'Rep Ventas', 'Día 0'],
    ['2. Auditoría 15 Min', 'Diagnóstico de 5 áreas + clasificación', 'Rep Ventas', 'Día 0–1'],
    ['3. Intake Completo', 'Formulario de intake, selección de paquete/tier', 'Rep Ventas', 'Día 1–2'],
    ['4. Pago y Contrato', 'Firma de contrato, cobro de Setup Fee', 'Rep Ventas + Fundador', 'Día 2–5'],
    ['5. Kickoff Técnico', 'Reunión de inicio, recolección de materiales', 'Técnico', 'Día 5–7'],
    ['6. Implementación', 'Construcción y configuración', 'Técnico', 'Según paquete'],
    ['7. Go-Live', 'Activación, capacitación, handoff a soporte', 'Técnico + Soporte', 'Día final'],
])

doc.add_heading('Etapa 1: Captura y Asignación del Lead', level=2)
ap('1. Rep registra el lead en el CRM con datos de contacto.')
ap('2. Clasificar tipo de negocio: Emprendedor, Comercio, Servicios, Hospedaje, Empresa.')
ap('3. Asignar prioridad: Alta (listo para decidir), Media (interesado), Baja (informativo).')
ap('4. Si el lead apunta a PRO AUTOMATION, escalar al Fundador.')

doc.add_heading('Etapa 2: Auditoría de 15 Minutos', level=2)
ap('OBLIGATORIO antes de recomendar paquete:', bold=True)
ap('1. Evaluar 5 áreas: Leads, Velocidad de Respuesta, Ventas, Operaciones, Retención.')
ap('2. Clasificar resultado:')
ap('   • No listo digitalmente → STARTER')
ap('   • Necesita bookings/leads + automatización → GROWTH')
ap('   • Dolor operacional + workflows complejos → PRO AUTOMATION')
ap('3. Documentar hallazgos en el formulario de intake.')

doc.add_heading('Etapa 3: Formulario de Intake Completo', level=2)
ap('1. Llenar formulario de intake con el prospecto.')
ap('2. Seleccionar paquete y tier específico basado en auditoría.')
ap('3. Identificar necesidades adicionales:')
ap('   • ¿Necesita módulo de reservas? (capacidad universal, no add-on)')
ap('   • ¿Necesita Marketing Digital? (add-on opcional)')
ap('   • ¿Necesita workflows adicionales? (bundles)')
ap('4. Verificar capacidad antes de avanzar.')

doc.add_heading('Etapa 4: Pago y Contrato', level=2)
ap('1. Presentar cotización formal con el paquete seleccionado.')
ap('2. Ofrecer opciones de pago:', bold=True)
ap('   • Mes a mes (estándar)')
ap('   • 6 meses upfront → 1 mes gratis o descuento setup')
ap('   • 12 meses upfront → 2 meses gratis o descuento setup')
ap('   • Plan Flex → Setup 50% hoy + 50% en 30 días')
ap('3. Enviar contrato de servicios para firma.')
ap('4. Cobrar Setup Fee (completo o 50% si Plan Flex).')
ap('5. Registrar pago en el dashboard financiero.')
ap('6. Confirmar fecha de inicio con el equipo técnico.')

doc.add_heading('Etapa 5: Kickoff Técnico', level=2)
ap('1. Programar reunión de kickoff dentro de 5 días hábiles post-pago.')
ap('2. Presentar el equipo al cliente.')
ap('3. Revisar alcance del servicio y entregables pactados.')
ap('4. Solicitar materiales necesarios: logo, contenido, accesos, credenciales.')
ap('5. Establecer canal de comunicación principal (WhatsApp Business o email).')
ap('6. Definir interlocutor del cliente con poder de decisión.')

doc.add_heading('Etapa 6: Implementación y Construcción', level=2)
ap('Tiempos de implementación por paquete:', bold=True)
at(['Paquete', 'Tiempo Estimado', 'Complejidad'], [
    ['STARTER ($9.99–$49.99)', '1–5 días según tier', 'Baja'],
    ['GROWTH ($50–$100)', '1–3 semanas', 'Media'],
    ['PRO AUTOMATION', '4–8 semanas (por fases)', 'Alta'],
    ['Add-On: Marketing Digital', '1 semana (setup campañas)', 'Media'],
])
ap('Durante la implementación:')
ap('• Actualizar progreso en el CRM cada 3 días')
ap('• Enviar actualización al cliente semanalmente')
ap('• Documentar cualquier cambio de alcance (requiere adenda formal)')
ap('• Si el cliente retrasa entrega de materiales, documentar y ajustar timeline')

doc.add_heading('Etapa 7: Go-Live y Activación', level=2)
ap('1. Activar todos los sistemas en producción.')
ap('2. Realizar pruebas finales de funcionamiento.')
ap('3. Capacitar al cliente en el uso de los sistemas.')
ap('4. Entregar documentación de usuario si aplica.')
ap('5. Hacer handoff formal al equipo de soporte.')
ap('6. Enviar mensaje de bienvenida post-lanzamiento.')
ap('7. Programar check-in de seguimiento a 2 semanas.')

# === SECTION 4 ===
doc.add_heading('SECCIÓN 4: CHECKLISTS DE ONBOARDING POR PAQUETE', level=1)

doc.add_heading('4.1 STARTER ($9.99–$49.99)', level=2)
ap('☐ Workflow Lite seleccionado del menú (5 opciones)')
ap('☐ Plantilla de workflow configurada y probada')
ap('☐ Canal de envío configurado (WhatsApp O email)')
ap('☐ Plantillas de contenido entregadas al cliente')
ap('☐ Google Business optimizado (tier $19.99+)')
ap('☐ Link hub page O booking link configurado (tier $29.99+)')
ap('☐ Micro-sitio 1 página creado (tier $49.99 solo)')
ap('☐ Cliente capacitado en uso del workflow')

doc.add_heading('4.2 GROWTH ($50–$100)', level=2)

doc.add_heading('Infraestructura Web', level=3)
ap('☐ Dominio registrado o conectado')
ap('☐ Hosting configurado')
ap('☐ Diseño del sitio web aprobado por el cliente')
ap('☐ Contenido recibido (textos, imágenes, logo)')
ap('☐ SEO básico configurado')
ap('☐ Google Business Profile configurado')
ap('☐ Google Analytics y Search Console conectados')
ap('☐ Captura de leads → notificación configurada')
ap('☐ CTA buttons y formularios integrados')

doc.add_heading('Módulo de Reservas (si aplica)', level=3)
ap('☐ Sistema de reservas configurado e integrado al sitio')
ap('☐ Calendario de disponibilidad configurado')
ap('☐ Notificaciones automáticas de reserva activadas')
ap('☐ Flujo completo de reserva probado')

doc.add_heading('Workflows', level=3)
ap('☐ Workflow Lite configurado (Growth $50)')
ap('☐ Full Workflow(s) configurado(s) y probado(s) (Growth $75/$100)')
ap('☐ Integraciones ligeras conectadas (Sheets/CRM/email/WhatsApp)')
ap('☐ Lógica if/then verificada (si aplica)')

doc.add_heading('4.3 PRO AUTOMATION', level=2)
ap('☐ Auditoría de procesos completada')
ap('☐ Mapa de procesos actuales documentado')
ap('☐ Puntos de dolor y cuellos de botella identificados')
ap('☐ Plan de implementación por fases aprobado')
ap('☐ 4+ workflows personalizados desarrollados')
ap('☐ Cada automatización probada end-to-end')
ap('☐ CRM avanzado configurado con campos personalizados')
ap('☐ Integraciones con sistemas existentes del cliente')
ap('☐ Reportería y dashboards de KPIs configurados')
ap('☐ Capacitación del equipo del cliente completada')
ap('☐ Documentación de procesos entregada')

doc.add_heading('4.4 Add-On: Marketing Digital', level=2)
ap('☐ Cuentas publicitarias del cliente configuradas (Meta, Google)')
ap('☐ Pixel/tracking instalado en sitio web')
ap('☐ Audiencias definidas')
ap('☐ Creativos de campaña aprobados por el cliente')
ap('☐ Campañas lanzadas y monitoreadas')
ap('☐ Integración con captura de leads verificada')
ap('☐ NOTA: Presupuesto de pauta es responsabilidad del cliente', italic=True)

# === SECTION 5 ===
doc.add_heading('SECCIÓN 5: REGLAS DE COMUNICACIÓN INTERNA', level=1)

doc.add_heading('Canales de Comunicación', level=2)
at(['Canal', 'Uso', 'Tiempo de Respuesta'], [
    ['WhatsApp (grupo equipo)', 'Actualizaciones rápidas, coordinación diaria', '< 2 horas'],
    ['Email', 'Comunicaciones formales, documentos, contratos', '< 24 horas'],
    ['CRM (notas)', 'Historial de cliente, actualizaciones de proyecto', 'Mismo día'],
    ['Reunión diaria (standup)', 'Revisión de pipeline, bloqueos, prioridades', 'Diaria 15 min'],
])

doc.add_heading('Reglas Generales', level=2)
ap('• No responder a clientes fuera de horario laboral (salvo emergencias)')
ap('• Todo compromiso con el cliente debe quedar por escrito')
ap('• Cambios de alcance requieren aprobación del Fundador')
ap('• Problemas técnicos graves: escalar inmediatamente')

# === SECTION 6 ===
doc.add_heading('SECCIÓN 6: PLANTILLAS DE COMUNICACIÓN CON EL CLIENTE', level=1)

doc.add_heading('6.1 Mensaje de Bienvenida (Post-Cierre)', level=2)
ap('¡Bienvenido/a a Cresca! 🎉', bold=True)
ap('Estamos emocionados de comenzar a trabajar contigo. Tu proyecto ya está en nuestro pipeline y dentro de los próximos [X] días hábiles te contactará nuestro equipo para comenzar.')
ap('¡Vamos a construir algo increíble juntos!')
ap('— Equipo Cresca')

doc.add_heading('6.2 Actualización de Progreso', level=2)
ap('Hola [Nombre],')
ap('Te compartimos una actualización de tu proyecto:')
ap('✅ Completado: [elementos terminados]')
ap('🔄 En progreso: [elementos activos]')
ap('⏳ Pendiente: [próximos pasos]')
ap('— Equipo Cresca')

doc.add_heading('6.3 Notificación de Go-Live', level=2)
ap('🚀 ¡Tu sistema está en vivo!', bold=True)
ap('Hola [Nombre], tu sistema ya está activo y funcionando.')
ap('Te contactaremos en 2 semanas para un check-in de seguimiento.')
ap('— Equipo Cresca')

doc.add_heading('6.4 Recordatorio de Pago', level=2)
ap('Hola [Nombre],')
ap('Recordatorio amigable de que tu cuota mensual de $[monto] venció el [fecha].')
ap('Si ya realizaste el pago, por favor ignora este mensaje.')
ap('— Equipo Cresca')

doc.save(r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\.tmp\sop_part1.docx')
print("Part 1 saved (Sections 1-6)")
