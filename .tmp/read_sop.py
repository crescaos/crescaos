import docx
import sys
sys.stdout.reconfigure(encoding='utf-8')

# SOP full content
doc = docx.Document('Todo_Directo_SOP_v2.docx')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)
for i, table in enumerate(doc.tables):
    print(f"\n--- Table {i+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(' | '.join(cells))
