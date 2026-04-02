"""Read full pricing strategy content."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

BASE = r'c:\Users\12132\Desktop\DigiFacil\DigiFacil'
doc = Document(os.path.join(BASE, 'Todo_Directo_Pricing_Strategy.docx'))

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[{p.style.name}] {t}")

for ti, table in enumerate(doc.tables):
    print(f"\n--- TABLE {ti+1} ---")
    for row in table.rows:
        print(" | ".join([c.text.strip() for c in row.cells]))
