"""Generate CRESCA ES Founding Rep Compensation Plan — Separate Document."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\CRESCA_ES_Founding_Rep_Compensation.docx'
doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sn = doc.styles['Normal']
sn.font.name = 'Calibri'; sn.font.size = Pt(11)
sn.paragraph_format.space_after = Pt(6)

for lvl in [1,2,3]:
    st = doc.styles[f'Heading {lvl}']
    st.font.name = 'Calibri'; st.font.bold = True
    st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    st.font.size = Pt([0,18,14,12][lvl])

def ap(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p

def at(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c = t.cell(0,j); c.text = ''; run = c.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.name = 'Calibri'
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = t.cell(i+1,j); c.text = ''; run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10); run.font.name = 'Calibri'
    doc.add_paragraph()
    return t

# ============ CONTENT ============

doc.add_heading('CRESCA — Plan de Compensación', level=1)
doc.add_heading('Representante Fundador · El Salvador', level=2)
ap('Sistemas de Crecimiento Empresarial', bold=True)
ap('Documento Confidencial — Versión 1.0 — Marzo 2026', italic=True)
ap('')
ap('Este documento define el plan de compensación EXCLUSIVO para el primer representante de ventas en El Salvador (Rep Fundador ES). Este plan es único y no aplica para afiliados, contractors independientes ni reps futuros.', bold=True)

# --- Section 1 ---
doc.add_heading('1. RESUMEN DE COMPENSACIÓN', level=2)

at(['Componente', 'Detalle'], [
    ['Salario Base', '$600 USD/mes (garantizado)'],
    ['Comisión de Setup', '20% de los fees de setup COBRADOS'],
    ['Comisión Residual', '6% del ingreso recurrente mensual (MRR) COBRADO'],
    ['Clawback de Calidad', '50% de la comisión de setup si cliente cancela dentro de 60 días por mal fit/oversell'],
])

# --- Section 2 ---
doc.add_heading('2. SALARIO BASE', level=2)
ap('Monto: $600 USD por mes.', bold=True)
ap('• Garantizado independientemente de resultados de ventas')
ap('• Se paga quincenalmente o mensualmente según acuerdo')
ap('• Sujeto a retenciones legales aplicables en El Salvador')
ap('• El salario base se revisa anualmente o cuando se alcancen hitos de crecimiento')

# --- Section 3 ---
doc.add_heading('3. COMISIÓN POR SETUP FEE (20%)', level=2)
ap('Cada vez que un cliente paga su fee de configuración inicial (Setup Fee), el representante recibe el 20% del monto COBRADO.', bold=True)

doc.add_heading('Ejemplos de Setup Commission', level=3)
at(['Paquete', 'Setup Fee', 'Tu Comisión (20%)'], [
    ['STARTER Assist', '$25', '$5.00'],
    ['STARTER Presence', '$75', '$15.00'],
    ['STARTER Pro', '$150', '$30.00'],
    ['GROWTH Foundation', '$250', '$50.00'],
    ['GROWTH +1 WF', '$350', '$70.00'],
    ['GROWTH +2 WF', '$450', '$90.00'],
    ['PRO AUTOMATION (ejemplo)', '$1,500', '$300.00'],
])

ap('Nota: STARTER DIY Lite ($9.99/mes) no tiene Setup Fee → no genera comisión de setup.', italic=True)

# --- Section 4 ---
doc.add_heading('4. COMISIÓN RESIDUAL (6%)', level=2)
ap('Cada mes que un cliente permanezca activo y al corriente en su pago, el representante recibe el 6% de la cuota mensual cobrada.', bold=True)

doc.add_heading('Ejemplos de Residual Commission', level=3)
at(['Paquete', 'Cuota Mensual', 'Tu Residual (6%/mes)'], [
    ['STARTER DIY Lite', '$9.99', '$0.60'],
    ['STARTER Assist', '$19.99', '$1.20'],
    ['STARTER Presence', '$29.99', '$1.80'],
    ['STARTER Pro', '$49.99', '$3.00'],
    ['GROWTH Foundation', '$50', '$3.00'],
    ['GROWTH +1 WF', '$75', '$4.50'],
    ['GROWTH +2 WF', '$100', '$6.00'],
    ['PRO AUTOMATION (ejemplo)', '$500', '$30.00'],
])

ap('IMPORTANTE: El residual es ACUMULATIVO. Cada nuevo cliente que cierres y permanezca activo incrementa tu ingreso mensual pasivo.', bold=True)

# --- Section 5 ---
doc.add_heading('5. PROTECCIÓN DE CALIDAD — CLAWBACK', level=2)
ap('Si un cliente cancela dentro de los primeros 60 días DEBIDO A mal fit o oversell (el representante vendió algo que no se ajustaba al negocio del cliente), la empresa recuperará el 50% de la comisión de setup de ese cliente.', bold=True)

doc.add_heading('Reglas del Clawback', level=3)
ap('• Solo aplica si la causa de cancelación es atribuible al representante')
ap('• Si la cancelación es por razones del cliente (presupuesto, cierre de negocio, etc.), NO aplica clawback')
ap('• El clawback se deduce del próximo pago de comisiones')
ap('• Si no hay comisiones pendientes, se deduce del siguiente mes')

doc.add_heading('Ejemplo de Clawback', level=3)
at(['Evento', 'Monto'], [
    ['Setup Fee del cliente cancelado', '$450 (GROWTH +2 WF)'],
    ['Tu comisión de ese setup (20%)', '$90'],
    ['Clawback (50% de tu comisión)', '$45 deducidos del próximo pago'],
])

# --- Section 6 ---
doc.add_heading('6. REGLAS DE PAGO — OBLIGATORIAS', level=2)

doc.add_heading('6.1 Solo sobre dinero cobrado', level=3)
ap('Las comisiones se calculan y pagan ÚNICAMENTE sobre dinero que la empresa ha cobrado efectivamente del cliente. Si un cliente firma contrato pero no paga, no hay comisión.', bold=True)

doc.add_heading('6.2 Residual solo mientras esté activo', level=3)
ap('El residual mensual se paga SOLO mientras el cliente esté activo y al corriente en sus pagos.')
ap('• Si un cliente pausa su cuenta → residual se pausa también')
ap('• Si un cliente cancela → residual deja de pagarse')
ap('• Si un cliente reactiva → residual se reanuda')

doc.add_heading('6.3 Exclusiones — Sin comisión sobre:', level=3)
ap('• Presupuesto de pauta publicitaria (ad spend)')
ap('• Fees de terceros (hosting, plataformas, herramientas)')
ap('• Reembolsos a clientes')
ap('• Regalías, descuentos u ofertas especiales')

# --- Section 7 ---
doc.add_heading('7. PROYECCIONES DE INGRESO', level=2)
ap('Supuestos: cierra 3 clientes/mes con mix de STARTER y GROWTH. Setup promedio ~$200, cuota promedio ~$55.', italic=True)

at(['Mes', 'Clientes\nActivos', 'Base', 'Setup\n(20%)', 'Residual\nAcumulado (6%)', 'Total\nMes'], [
    ['1', '3', '$600', '$120', '$10', '$730'],
    ['3', '9', '$600', '$120', '$30', '$750'],
    ['6', '18', '$600', '$120', '$59', '$779'],
    ['9', '27', '$600', '$120', '$89', '$809'],
    ['12', '36', '$600', '$120', '$119', '$839'],
])

ap('Escenario más agresivo (5 clientes/mes, mix con más GROWTH):', bold=True)
at(['Mes', 'Clientes\nActivos', 'Base', 'Setup\n(20%)', 'Residual\nAcumulado (6%)', 'Total\nMes'], [
    ['1', '5', '$600', '$200', '$17', '$817'],
    ['6', '30', '$600', '$200', '$99', '$899'],
    ['12', '60', '$600', '$200', '$198', '$998'],
])

ap('A 12 meses con buen rendimiento, tu ingreso mensual puede acercarse a los $1,000 USD entre base, setup y residual acumulado. El residual seguirá creciendo cada mes.', bold=True)

# --- Section 8 ---
doc.add_heading('8. CONDICIONES ESPECIALES', level=2)

ap('• Este plan de compensación es para el ROL de Representante Fundador ES únicamente')
ap('• No aplica para afiliados, referidos ni contractors independientes')
ap('• La empresa se reserva el derecho de revisar este plan con 30 días de aviso')
ap('• Cualquier modificación será por escrito y firmada por ambas partes')
ap('• Los datos de clientes, contactos y procesos de Cresca son confidenciales')

# --- Section 9 ---
doc.add_heading('9. RESUMEN EJECUTIVO', level=2)

at(['Elemento', 'Detalle'], [
    ['Rol', 'Representante Fundador — El Salvador'],
    ['Salario Base', '$600 USD/mes (garantizado)'],
    ['Comisión Setup', '20% del fee cobrado'],
    ['Comisión Residual', '6% del MRR cobrado (acumulativo)'],
    ['Clawback', '50% de comisión setup si cancela en 60 días por mal fit'],
    ['Pago', 'Solo sobre dinero cobrado, no facturado'],
    ['Exclusiones', 'Sin comisión sobre ad spend, fees 3ros, reembolsos'],
    ['Residual', 'Solo mientras cliente esté activo y al corriente'],
])

ap('')
ap('Cresca — Sistemas de Crecimiento Empresarial', bold=True)
ap('Plan de Compensación — Representante Fundador ES', italic=True)
ap('Documento confidencial para uso interno. Versión 1.0 — Marzo 2026', italic=True)

doc.save(OUT)
print(f"Saved: {OUT}")
print("Done!")
