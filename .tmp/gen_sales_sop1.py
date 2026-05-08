"""Generate Cresca Sales SOP Part 1 (Sections 1-7) — 3-Tier Model."""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

sn = doc.styles['Normal']; sn.font.name = 'Calibri'; sn.font.size = Pt(11)
sn.paragraph_format.space_after = Pt(4)
for lvl in [1,2,3]:
    st = doc.styles[f'Heading {lvl}']
    st.font.name = 'Calibri'; st.font.bold = True
    st.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    st.font.size = Pt([0,18,14,12][lvl])

def ap(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p

def at(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c = t.cell(0,j); c.text = ''; run = c.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.name = 'Calibri'
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = t.cell(i+1,j); c.text = ''; run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10); run.font.name = 'Calibri'
    doc.add_paragraph(); return t

# ============ CONTENT ============

doc.add_heading('CRESCA — Manual de Ventas (SOP)', level=1)
ap('Sistemas de Crecimiento Empresarial', bold=True)
ap('Guía Completa para Representantes de Ventas', italic=True)
ap('Versión 3.0 — Marzo 2026', italic=True)

# TOC
doc.add_heading('TABLA DE CONTENIDOS', level=1)
for s in [
    'Sección 1: Tu Portafolio — Tres Paquetes, Una Misión',
    'Sección 2: La Oportunidad en El Salvador',
    'Sección 3: Sistema de Venta Consultiva (SPIN)',
    'Sección 4: Árbol de Decisión — Auditoría en 15 Minutos',
    'Sección 5: El Proceso Universal de Ventas en 6 Pasos',
    'Sección 6: Scripts Detallados por Paquete',
    'Sección 7: Workflow Lite, Full Workflows y Add-Ons',
    'Sección 8: Cresca Como Tu Diferenciador',
    'Sección 9: Guía de Manejo de Objeciones',
    'Sección 10: Estructura de Compensación',
    'Sección 11: Herramientas y Recursos de Ventas',
    'Sección 12: Guía de Campo — Flujo Diario',
    'Sección 13: Tarjeta de Referencia Rápida',
]:
    ap(s)

# === SECTION 1 ===
doc.add_heading('SECCIÓN 1: TU PORTAFOLIO — TRES PAQUETES, UNA MISIÓN', level=1)
ap('Cresca vende sistemas de crecimiento empresarial, organizados en una escalera de 3 niveles.', bold=True)
ap('Tu trabajo como representante es diagnosticar qué nivel necesita cada negocio y guiarlos hacia la solución correcta usando la auditoría de 15 minutos.')

doc.add_heading('Paquete 1: STARTER (Micro/Pequeño)', level=2)
at(['Característica', 'Detalle'], [
    ['¿Qué es?', 'Entrada al mundo digital con herramientas básicas y 1 Workflow Lite incluido'],
    ['Rango mensual', '$9.99 – $49.99'],
    ['Setup Fee', '$0 – $199 según tier'],
    ['Tiers', 'DIY Lite ($9.99) → Starter Assist ($19.99) → Starter Presence ($29.99) → Starter Pro ($49.99)'],
    ['Incluye', '1 Workflow Lite (plantilla, 1 canal, sin ramificación)'],
    ['Para quién', 'Negocios que no están listos digitalmente y necesitan empezar desde lo básico'],
    ['Pitch en 1 frase', '"Empiece su presencia digital por menos de lo que cuesta un almuerzo al día."'],
])

doc.add_heading('Paquete 2: GROWTH (Infraestructura)', level=2)
at(['Característica', 'Detalle'], [
    ['¿Qué es?', 'Infraestructura digital completa con workflows de automatización'],
    ['Rango mensual', '$50 – $100'],
    ['Setup Fee', '$199 – $499 según tier'],
    ['Tiers', 'Foundation ($50) → Foundation +1 ($75) → Foundation +2 ($100)'],
    ['Incluye', '1 Workflow Lite a 2 Full Workflows según tier'],
    ['Para quién', 'Negocios que necesitan reservas/leads + automatización simple'],
    ['Pitch en 1 frase', '"Construimos tu infraestructura digital para que captures clientes y operes automáticamente."'],
])

doc.add_heading('Paquete 3: PRO AUTOMATION (Avanzado)', level=2)
at(['Característica', 'Detalle'], [
    ['¿Qué es?', 'Automatización avanzada de procesos y operaciones internas'],
    ['Precio', 'Cotización personalizada'],
    ['Setup Fee', 'Incluido en cotización'],
    ['Incluye', 'Auditoría + mapeo + 4+ workflows + monitoreo + optimización'],
    ['Para quién', 'Negocios con dolor operacional y procesos manuales que necesitan workflows complejos'],
    ['Contrato recomendado', '12 meses'],
    ['Pitch en 1 frase', '"Automatizamos tus procesos internos para que tu negocio crezca sin necesitar más empleados."'],
])

doc.add_heading('Resumen del Portafolio', level=3)
at(['Paquete', 'Precio Mensual', 'Setup Fee', 'Workflows Incluidos'], [
    ['STARTER', '$9.99 – $49.99', '$0 – $199', '1 Workflow Lite'],
    ['GROWTH', '$50 – $100', '$199 – $499', '1 Lite a 2 Full'],
    ['PRO AUTOMATION', 'Cotización', 'En cotización', '4+ personalizados'],
    ['+ Marketing Digital (Add-On)', 'Variable', 'Variable', '—'],
])

# === SECTION 2 ===
doc.add_heading('SECCIÓN 2: LA OPORTUNIDAD EN EL SALVADOR', level=1)

doc.add_heading('El Mercado Actual', level=2)
ap('La mayoría de los negocios en El Salvador operan con:')
ap('• Procesos 100% manuales (cuadernos, WhatsApp personal, hojas de Excel)')
ap('• Sin presencia web profesional')
ap('• Dependencia total de Facebook Marketplace o Instagram para ventas')
ap('• Sin CRM — pierden clientes por falta de seguimiento')
ap('• Sin automatización — el dueño hace todo personalmente')

doc.add_heading('Negocios Objetivo por Paquete', level=2)
at(['Tipo de Negocio', 'Paquete Ideal', 'Por Qué'], [
    ['Emprendedor iniciando', 'STARTER', 'Necesita empezar desde cero, bajo presupuesto'],
    ['Restaurantes / Cafés', 'GROWTH', 'Necesitan web, reservas, captura de leads'],
    ['Tiendas / Comercio', 'STARTER o GROWTH', 'Según madurez digital y presupuesto'],
    ['Servicios profesionales', 'GROWTH', 'Necesitan reservas, seguimiento, presencia web'],
    ['Clínicas / Salud', 'GROWTH o PRO AUTOMATION', 'Procesos complejos, citas, seguimiento'],
    ['Hospedaje / Inmobiliarias', 'GROWTH (con módulo reservas)', 'Necesitan reservas directas'],
    ['Empresas +10 empleados', 'PRO AUTOMATION', 'Procesos manuales, ineficiencia operativa'],
])

doc.add_heading('Por Qué Los Sistemas Importan', level=2)
ap('Un negocio sin sistemas es un negocio que depende 100% del dueño.', bold=True)
ap('Cresca construye sistemas que funcionan 24/7:')
ap('• El sitio web captura clientes mientras el dueño duerme')
ap('• El workflow responde y da seguimiento automáticamente')
ap('• El CRM hace seguimiento automático sin que nadie lo recuerde')
ap('• Las automatizaciones ahorran horas de trabajo manual cada día')

# === SECTION 3 ===
doc.add_heading('SECCIÓN 3: SISTEMA DE VENTA CONSULTIVA (SPIN)', level=1)
ap('Cresca usa el método SPIN para venta consultiva. NO vendas — descubre el problema y conecta con la solución.', bold=True)

doc.add_heading('1. Pregunta de Situación', level=2)
ap('Objetivo: Entender el estado actual del negocio.', italic=True)
ap('• "¿Cómo manejas tus clientes actualmente?"')
ap('• "¿Tienes sitio web o algún sistema digital?"')
ap('• "¿Cuántos empleados tienes y qué herramientas usan?"')
ap('• "¿Cómo te encuentran tus clientes nuevos?"')

doc.add_heading('2. Pregunta de Problema', level=2)
ap('Objetivo: Identificar dolores y frustraciones.', italic=True)
ap('• "¿Sientes que pierdes clientes por no responder a tiempo?"')
ap('• "¿Cuántas horas al día dedicas a tareas administrativas?"')
ap('• "¿Te ha pasado que un cliente se fue a la competencia porque no le diste seguimiento?"')

doc.add_heading('3. Pregunta de Implicación', level=2)
ap('Objetivo: Hacer que sientan el costo de NO resolver el problema.', italic=True)
ap('• "Si sigues manejando todo por WhatsApp personal, ¿cuántos clientes más vas a perder este año?"')
ap('• "¿Qué pasa con tu negocio cuando tú te enfermas o sales de viaje?"')

doc.add_heading('4. Pregunta de Necesidad-Beneficio', level=2)
ap('Objetivo: Que ellos mismos reconozcan que necesitan una solución.', italic=True)
ap('• "¿Cómo cambiaría tu negocio si tuvieras un sistema que responde clientes automáticamente?"')
ap('• "¿Qué harías con 20 horas extra al mes si tus procesos fueran automáticos?"')

doc.add_heading('5. Pregunta de Calificación', level=2)
ap('Objetivo: Determinar qué paquete recomendar.', italic=True)
ap('Si no está listo digitalmente, necesita empezar:')
ap('→ STARTER ($9.99–$49.99)', bold=True)
ap('')
ap('Si necesita reservas/leads + automatización simple:')
ap('→ GROWTH ($50–$100)', bold=True)
ap('')
ap('Si tiene dolor operacional y procesos manuales complejos:')
ap('→ PRO AUTOMATION (Cotización)', bold=True)

# === SECTION 4 ===
doc.add_heading('SECCIÓN 4: ÁRBOL DE DECISIÓN — AUDITORÍA EN 15 MINUTOS', level=1)
ap('OBLIGATORIO: El representante DEBE hacer una auditoría rápida de 15 minutos para diagnosticar necesidades antes de recomendar un paquete.', bold=True)

doc.add_heading('Las 5 Áreas de la Auditoría', level=2)
at(['Área', 'Pregunta Clave', 'Qué Evalúas'], [
    ['1. Leads', '¿De dónde vienen sus clientes potenciales?', 'Fuentes de leads, presencia digital actual'],
    ['2. Velocidad de Respuesta', '¿Cuánto tardan en responder a un prospecto?', 'Tiempo de respuesta, canales usados'],
    ['3. Reservas/Ventas', '¿Cómo manejan el proceso de cierre?', 'Proceso de venta, sistema de booking'],
    ['4. Operaciones', '¿Qué procesos internos consumen más tiempo?', 'Tareas repetitivas, ineficiencias'],
    ['5. Retención', '¿Cómo mantienen contacto con clientes existentes?', 'Follow-up, reseñas, fidelización'],
])

doc.add_heading('Regla de Decisión Post-Auditoría', level=2)
at(['Diagnóstico', 'Paquete', 'Acción'], [
    ['No está listo digitalmente', 'STARTER', 'Vender tier apropiado ($9.99–$49.99)'],
    ['Necesita bookings/leads + automatización simple', 'GROWTH', 'Vender Foundation, +1 o +2 ($50–$100)'],
    ['Dolor operacional + workflows complejos', 'PRO AUTOMATION', 'Programar auditoría completa y cotizar'],
])

doc.add_heading('PASO 2: Verificar Necesidades Adicionales', level=2)
ap('¿Necesita módulo de reservas? → Incluir en Growth o Pro Automation')
ap('¿Necesita publicidad digital? → Agregar Add-On Marketing Digital')
ap('¿Necesita más workflows? → Ofrecer bundles de workflows adicionales')

# === SECTION 5 ===
doc.add_heading('SECCIÓN 5: EL PROCESO UNIVERSAL DE VENTAS EN 6 PASOS', level=1)

doc.add_heading('PASO 1: APERTURA / ACERCAMIENTO', level=2)
ap('Puerta fría / WhatsApp / Referido — cualquier canal.', italic=True)
ap('Script de apertura:', bold=True)
ap('"Hola, [nombre]. Soy [tu nombre] de Cresca — Sistemas de Crecimiento Empresarial. Ayudamos a negocios como el tuyo a capturar más clientes y operar más eficientemente con sistemas digitales. ¿Tienes 15 minutos para una auditoría gratuita de tu negocio?"')
ap('')
ap('Lo que NO decir:', bold=True)
ap('• "Hacemos páginas web" (suena genérico)')
ap('• "Somos una agencia de marketing" (NO somos eso)')
ap('• "Vendemos chatbots" (NO vendemos productos sueltos)')

doc.add_heading('PASO 2: AUDITORÍA DE 15 MINUTOS', level=2)
ap('Usa la auditoría de 5 áreas (Sección 4). Escucha más de lo que hablas.', italic=True)
ap('Objetivo: Diagnosticar al prospecto y determinar:', bold=True)
ap('• ¿Está listo digitalmente o necesita empezar desde cero?')
ap('• ¿Necesita infraestructura o automatización avanzada?')
ap('• ¿Tiene presupuesto para invertir?')
ap('• ¿Quién toma las decisiones?')

doc.add_heading('PASO 3: PRESENTAR LA SOLUCIÓN', level=2)
ap('Conecta los hallazgos de la auditoría con el paquete correcto:', italic=True)
ap('"Basándome en lo que encontramos en la auditoría, lo que tu negocio necesita es [STARTER / GROWTH / PRO AUTOMATION]. Déjame explicarte exactamente qué incluye."')

doc.add_heading('PASO 4: DEMOSTRACIÓN', level=2)
ap('Muestra ejemplos reales de lo que Cresca construye:', italic=True)
ap('• Para STARTER: plantillas de WhatsApp, workflow lite en acción')
ap('• Para GROWTH: sitio web de un cliente, sistema de reservas, workflows')
ap('• Para PRO AUTOMATION: dashboard de reportes, automatizaciones complejas')

doc.add_heading('PASO 5: COTIZACIÓN Y CIERRE', level=2)
ap('Presenta la cotización con confianza:', italic=True)
ap('"La inversión para tu negocio es de $[Setup Fee] de configuración inicial y $[Cuota Mensual] mensuales. Esto incluye todo lo que te mostré, más soporte continuo."')
ap('')
ap('Técnicas de cierre:', bold=True)
ap('• Cierre directo: "¿Procedemos con la configuración esta semana?"')
ap('• Cierre alternativo: "¿Prefieres empezar con Starter o ir directo a Growth?"')
ap('• Cierre de urgencia: "Solo tenemos capacidad para [X] proyectos más este mes."')
ap('')
ap('Opciones de pago:', bold=True)
ap('• Mes a mes (estándar)')
ap('• 6 meses por adelantado → 1 mes gratis')
ap('• 12 meses por adelantado → 2 meses gratis')
ap('• Plan Flex: Setup dividido 50% hoy + 50% en 30 días')

doc.add_heading('PASO 6: ENTREGA Y BIENVENIDA', level=2)
ap('Una vez cerrada la venta:')
ap('1. Enviar contrato de servicios para firma')
ap('2. Cobrar Setup Fee (completo o 50% si es Flex)')
ap('3. Registrar en el CRM con datos completos')
ap('4. Enviar formulario de intake')
ap('5. Avisar al equipo técnico para iniciar onboarding')
ap('6. Enviar mensaje de bienvenida al cliente')

at(['Paquete', 'Tiempo de Implementación'], [
    ['STARTER ($9.99–$49.99)', '1–5 días según tier'],
    ['GROWTH ($50–$100)', '1–3 semanas'],
    ['PRO AUTOMATION', '4–8 semanas (por fases)'],
    ['Add-On Marketing', '1 semana (setup de campañas)'],
])

# === SECTION 6 ===
doc.add_heading('SECCIÓN 6: SCRIPTS DETALLADOS POR PAQUETE', level=1)

doc.add_heading('PAQUETE STARTER: Scripts de Venta', level=2)
ap('Script de Presentación:', bold=True)
ap('"[Nombre], veo que tu negocio está empezando a crecer pero todavía no tiene una base digital sólida. Empiece con lo más básico — por menos de $10 al mes tiene un workflow automatizado que responde a sus clientes cuando usted no puede, más todas las plantillas y guías para construir su presencia digital."')
ap('"Y lo mejor es que cuando esté listo para crecer, puede subir al siguiente nivel sin empezar desde cero."')
ap('')
ap('Funciones clave que destacar:', bold=True)
ap('• Workflow Lite automatizado desde el primer día')
ap('• Plantillas profesionales listas para usar')
ap('• Sin compromiso largo — mes a mes desde $9.99')
ap('• Escalable — cuando crezca, simplemente sube de nivel')

doc.add_heading('PAQUETE GROWTH: Scripts de Venta', level=2)
ap('Script de Presentación:', bold=True)
ap('"[Nombre], tu negocio ya está funcionando, pero estás perdiendo clientes porque no tienes una infraestructura digital. Lo que hacemos en Cresca es construir todo junto: tu sitio web, tu sistema de captura de leads, tus workflows de automatización — todo integrado para que funcione 24/7."')
ap('"Por $50 al mes tienes la base completa. Si necesitas más automatización, subimos a $75 o $100 y añadimos workflows completos."')
ap('')
ap('Funciones clave que destacar:', bold=True)
ap('• Sitio web profesional que captura leads automáticamente')
ap('• Módulo de reservas (si aplica)')
ap('• Workflows que hacen seguimiento sin que tú hagas nada')
ap('• Google Business integrado para aparecer en búsquedas')
ap('• Todo integrado — no son herramientas sueltas, es un sistema')

doc.add_heading('PAQUETE PRO AUTOMATION: Scripts de Venta', level=2)
ap('Script de Presentación:', bold=True)
ap('"[Nombre], tu negocio ha crecido, pero noto que estás llegando a un punto donde los procesos manuales te están frenando. Lo que hacemos en Cresca es auditar tus procesos, identificar dónde estás perdiendo tiempo y dinero, y construir automatizaciones específicas para tu negocio."')
ap('"Piensa en ello como contratar un equipo digital que nunca se enferma, no pide vacaciones, y trabaja a la velocidad de la luz."')
ap('')
ap('Funciones clave que destacar:', bold=True)
ap('• Auditoría que identifica exactamente dónde se pierde dinero')
ap('• 4+ automatizaciones personalizadas para el negocio')
ap('• Reportes que se generan solos con datos en tiempo real')
ap('• CRM avanzado que conecta todos los canales')
ap('• Procesos que antes tomaban horas, ahora toman minutos')

# === SECTION 7 ===
doc.add_heading('SECCIÓN 7: WORKFLOW LITE, FULL WORKFLOWS Y ADD-ONS', level=1)

doc.add_heading('Workflow Lite (Incluido en STARTER)', level=2)
ap('Cada plan STARTER incluye 1 Workflow Lite. Menú de opciones:')
ap('1. Auto respuesta inicial (WhatsApp)')
ap('2. Mensaje de confirmación (al agendar/solicitar)')
ap('3. Recordatorio simple')
ap('4. Solicitud de reseña post-servicio')
ap('5. Seguimiento 1 paso (24h después)')
ap('')
ap('Reglas: Solo plantilla, 1 canal, 1 update/mes, sin integraciones complejas.', italic=True)
ap('Si el cliente quiere más → el upgrade es a GROWTH, NO se venden Lites adicionales.', bold=True)

doc.add_heading('Full Workflows (Incluidos en GROWTH)', level=2)
ap('Growth incluye hasta 2 Full Workflows según tier seleccionado.')
ap('Cada Full Workflow incluye: trigger + acción + 1–2 pasos, integración ligera, lógica if/then limitada.')

doc.add_heading('Workflows Adicionales (Add-On para GROWTH)', level=2)
at(['Bundle', 'Precio Mensual', 'Cuándo Proponerlo'], [
    ['+1 workflow', '+$15/mes', 'Cliente necesita 1 automatización extra'],
    ['+2 workflows', '+$25/mes', 'Más de una necesidad identificada'],
    ['+5 workflows', '+$60/mes', 'Negocio con múltiples procesos'],
    ['+8 workflows', '+$90/mes', 'Pre-Pro Automation, máximo en Growth'],
])

doc.add_heading('Add-On: Marketing Digital', level=2)
ap('• Gestión de campañas Meta/Google Ads')
ap('• Ad spend lo paga el cliente (separado)')
ap('• Solo se vende junto con GROWTH o PRO AUTOMATION')
ap('Script: "Ahora que tienes la infraestructura para capturar clientes, podemos activar campañas publicitarias para traer tráfico. Es como abrir la llave del agua ahora que tienes la tubería lista."')

ap('REGLA: Nunca vendas Marketing Digital solo ni con STARTER. El cliente debe tener al menos GROWTH para que las campañas tengan dónde aterrizar.', bold=True)

# Save part 1
doc.save(r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\.tmp\sales_sop_p1.docx')
print("Sales SOP Part 1 saved (Sections 1-7)")
