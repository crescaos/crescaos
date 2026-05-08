# Part 2: Sections 5-12 — loads sop_partial.docx and adds remaining sections
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document(r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\.tmp\sop_partial.docx')

HEADER_BG = '2E86C1'

def set_cell_bg(cell, color):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(s)

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255,255,255)
                r.font.size = Pt(10)
        set_cell_bg(c, HEADER_BG)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for ri in range(len(t.rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Inches(w)
    doc.add_paragraph()
    return t

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level*0.25)
    return p

# ========== SECTION 5 ==========
doc.add_page_break()
doc.add_heading('SECCIÓN 5: REGLAS DE COMUNICACIÓN INTERNA', level=1)
doc.add_paragraph(
    'La comunicación efectiva es la base de operaciones eficientes. Estas reglas definen '
    'cómo se comunica el equipo internamente para evitar malentendidos, retrasos y tareas perdidas.'
)

doc.add_heading('Canales de Comunicación', level=2)
add_table(doc, ['Canal', 'Uso', 'Tiempo de Respuesta', 'Quién Participa'],
    [
        ['WhatsApp – Grupo Operaciones', 'Actualizaciones rápidas, preguntas urgentes del día a día', '≤ 30 minutos', 'Todo el equipo'],
        ['WhatsApp – Canal de Cierres', 'Notificación de cada nueva venta cerrada por reps', 'Solo lectura', 'Reps + Ops'],
        ['Email (info@tododirecto.com)', 'Formularios de intake, archivos de clientes, documentación formal', '≤ 4 horas hábiles', 'Ops + Técnico'],
        ['Reunión Diaria (Standup)', 'Revisión de pipeline, bloqueos, prioridades del día', 'Diario 8:30 AM, 15 min', 'Ops + Técnico'],
        ['Reunión Semanal', 'Revisión de KPIs, pipeline completo, planificación de capacidad', 'Lunes 9:00 AM, 30 min', 'Todo el equipo'],
        ['[PLACEHOLDER] – Herramienta de Gestión', 'Tracking de proyectos, estados, asignaciones', 'Actualizar 2x al día', 'Ops + Técnico'],
    ],
    [2.0, 2.5, 1.5, 1.5])

doc.add_heading('Reglas de la Reunión Diaria (Standup)', level=2)
doc.add_paragraph('Formato fijo — cada persona responde 3 preguntas:')
add_bullet(doc, '¿Qué completé ayer?')
add_bullet(doc, '¿Qué haré hoy?')
add_bullet(doc, '¿Tengo algún bloqueo?')
p = doc.add_paragraph()
r = p.add_run('Regla: Si no hay bloqueos y no hay novedades, se dice "Sin novedades" — no se extiende.')
r.italic = True

doc.add_heading('Reglas Generales', level=2)
add_bullet(doc, 'Nunca comunicar un problema sin una solución propuesta.')
add_bullet(doc, 'Toda decisión que afecte al cliente debe documentarse por escrito (no solo verbal).')
add_bullet(doc, 'Si un proyecto está bloqueado más de 24 horas, escalar automáticamente al Gerente de Ops.')
add_bullet(doc, 'Los mensajes urgentes se marcan con "🔴 URGENTE:" al inicio.')
add_bullet(doc, 'Los mensajes que requieren acción se marcan con "📌 ACCIÓN:" al inicio.')

# ========== SECTION 6 ==========
doc.add_page_break()
doc.add_heading('SECCIÓN 6: PLANTILLAS DE COMUNICACIÓN CON EL CLIENTE', level=1)
doc.add_paragraph(
    'Todas las comunicaciones con clientes deben ser profesionales, consistentes y en la voz '
    'de Todo Directo. A continuación, las plantillas para los momentos clave del ciclo de vida del cliente.'
)

templates = [
    ('6.1 Mensaje de Bienvenida (Post-Cierre)', 
     '¡Hola [NOMBRE]! 🎉 Soy [NOMBRE_OPS] del equipo de operaciones de Todo Directo. Tu representante '
     '[NOMBRE_REP] nos ha informado que has contratado [PRODUCTO/PLAN]. ¡Bienvenido/a!\n\n'
     'Estos son los próximos pasos:\n'
     '1. Revisaremos tu formulario de orden de servicio.\n'
     '2. Te contactaremos en las próximas 48 horas para confirmar los detalles de tu proyecto.\n'
     '3. Tiempo estimado de entrega: [PLAZO].\n\n'
     '¿Alguna pregunta? Estoy a tu disposición. 😊'),
    ('6.2 Solicitud de Materiales Faltantes',
     'Hola [NOMBRE], espero que estés bien. Para avanzar con tu [PRODUCTO], necesitamos los '
     'siguientes materiales:\n\n'
     '📎 [LISTA DE MATERIALES FALTANTES]\n\n'
     'Puedes enviarlos por WhatsApp, Google Drive o correo a info@tododirecto.com.\n\n'
     '⏱ Nota: Tu proyecto está en pausa hasta que recibamos estos materiales. Una vez los tengamos, '
     'retomamos inmediatamente.'),
    ('6.3 Actualización de Progreso',
     'Hola [NOMBRE], actualización rápida sobre tu [PRODUCTO]:\n\n'
     '✅ Lo que ya completamos: [AVANCES]\n'
     '🔄 Lo que estamos trabajando: [EN PROGRESO]\n'
     '📅 Fecha estimada de revisión contigo: [FECHA]\n\n'
     'Te mantendremos informado/a. ¡Va quedando increíble! 🚀'),
    ('6.4 Invitación a Revisión',
     'Hola [NOMBRE], ¡tu [PRODUCTO] está listo para revisión! 🎯\n\n'
     '🔗 [ENLACE O INSTRUCCIONES PARA VER]\n\n'
     'Por favor revísalo y dinos:\n'
     '✅ Si todo está correcto y podemos activarlo.\n'
     '✏️ Si tienes cambios (incluye los detalles específicos).\n\n'
     '⏱ Tienes 48 horas para darnos feedback. Si no recibimos respuesta, asumiremos que está aprobado '
     'y procederemos con el go-live.'),
    ('6.5 Notificación de Go-Live',
     '¡Hola [NOMBRE]! 🎉🚀\n\n'
     'Tu [PRODUCTO] ya está ACTIVO y funcionando.\n\n'
     '🔗 Acceso: [ENLACE/INSTRUCCIONES]\n'
     '📱 App/Panel: [INSTRUCCIONES DE ACCESO]\n\n'
     'Puntos importantes:\n'
     '• Tu suscripción mensual de $[PRECIO] inicia hoy.\n'
     '• Te contactaremos en 2 semanas para un check-in.\n'
     '• Si tienes cualquier problema, escríbenos aquí.\n\n'
     '¡Gracias por confiar en Todo Directo! 💜'),
    ('6.6 Check-in Post-Lanzamiento (2 Semanas)',
     'Hola [NOMBRE], ¿cómo va todo con tu [PRODUCTO]? Ya llevas 2 semanas activo/a.\n\n'
     'Nos gustaría saber:\n'
     '1. ¿Está funcionando como esperabas?\n'
     '2. ¿Has tenido algún problema técnico?\n'
     '3. ¿Hay algo que podamos mejorar?\n\n'
     'Tu éxito es nuestra prioridad. 🎯'),
    ('6.7 Recordatorio de Pago',
     'Hola [NOMBRE], este es un recordatorio amistoso de que tu pago mensual de $[MONTO] por '
     '[PRODUCTO] vence el [FECHA].\n\n'
     'Métodos de pago:\n'
     '💳 Transferencia: [PLACEHOLDER – datos bancarios]\n'
     '📱 Chivo Wallet: [PLACEHOLDER]\n'
     '💵 Efectivo: coordinar con tu representante\n\n'
     'Si ya realizaste el pago, puedes ignorar este mensaje. ¡Gracias! 🙏'),
]

for title, body in templates:
    doc.add_heading(title, level=2)
    # Add as a bordered paragraph (using a 1-cell table for visual effect)
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    cell.text = body
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string('2C3E50')
    doc.add_paragraph()

# ========== SECTION 7 ==========
doc.add_page_break()
doc.add_heading('SECCIÓN 7: REGLAS DE PLANIFICACIÓN DE CAPACIDAD', level=1)
doc.add_paragraph(
    'Para mantener calidad y cumplir SLAs, el equipo debe operar dentro de límites de capacidad '
    'definidos. Si la demanda supera la capacidad, se activan las reglas de expansión.'
)

doc.add_heading('Límites de Capacidad por Rol', level=2)
add_table(doc, ['Rol', 'Proyectos Simultáneos Máx.', 'Tipo de Proyecto', 'Notas'],
    [
        ['Desarrollador Web', '3–4', 'Sitios web, apps', 'Depende de complejidad del plan'],
        ['Especialista IA/Chatbot', '2–3', 'Chatbots, automatizaciones', 'Mayor complejidad técnica'],
        ['QA / Soporte', '5–6', 'Revisiones, soporte', 'Tareas más cortas, mayor volumen'],
        ['Gerente de Operaciones', '10–12', 'Supervisión de pipeline', 'No implementa, coordina'],
    ],
    [2.0, 1.5, 1.5, 2.5])

doc.add_heading('Reglas de Expansión de Capacidad', level=2)
add_bullet(doc, 'Señal verde (< 70% capacidad): Operación normal. Aceptar todos los cierres.')
add_bullet(doc, 'Señal amarilla (70–90% capacidad): Priorizar proyectos por valor de contrato. Informar a reps que los plazos pueden extenderse 2–3 días.')
add_bullet(doc, 'Señal roja (> 90% capacidad): Escalar al CEO. Evaluar contratación de freelancer o pausa temporal de ventas de productos de alta complejidad.')

doc.add_heading('Calendario de Demanda Estacional', level=2)
add_table(doc, ['Período', 'Demanda Esperada', 'Productos Más Demandados', 'Acción'],
    [
        ['Ene–Feb', 'Media', 'Web, AlToque', 'Inicio de año — negocios renuevan'],
        ['Mar–Abr', 'Alta', 'Vacation Rentals, Web', 'Pre-Semana Santa turismo'],
        ['May–Jun', 'Media', 'Todos', 'Operación normal'],
        ['Jul–Ago', 'Alta', 'Vacation Rentals, Chatbot', 'Temporada alta de turismo'],
        ['Sep–Oct', 'Baja', 'Automatización', 'Temporada baja — foco en retención'],
        ['Nov–Dic', 'Alta', 'Todos', 'Black Friday, fiestas, turismo decembrino'],
    ],
    [1.2, 1.3, 2.0, 3.0])

# ========== SECTION 8 ==========
doc.add_page_break()
doc.add_heading('SECCIÓN 8: PROTOCOLO DE ESCALAMIENTO', level=1)
doc.add_paragraph(
    'No todos los problemas se resuelven al mismo nivel. Este protocolo define 3 niveles de '
    'escalamiento con tiempos de respuesta claros.'
)

doc.add_heading('Niveles de Escalamiento', level=2)
add_table(doc, ['Nivel', 'Responsable', 'Tipo de Problema', 'Tiempo de Respuesta', 'Ejemplo'],
    [
        ['Nivel 1', 'Soporte Técnico', 'Problemas técnicos menores, preguntas de uso, ajustes de contenido',
         '≤ 4 horas hábiles', 'Cliente no puede subir una foto; necesita cambiar un texto'],
        ['Nivel 2', 'Gerente de Operaciones', 'Problemas que afectan la entrega, disputas de servicio, bloqueos técnicos complejos',
         '≤ 8 horas hábiles', 'Integración no funciona; cliente insatisfecho con la implementación'],
        ['Nivel 3', 'CEO / Fundador', 'Amenaza de cancelación, problema legal, fallo de seguridad, crisis de reputación',
         '≤ 2 horas', 'Cliente amenaza con denuncia pública; fuga de datos'],
    ],
    [0.6, 1.2, 2.0, 1.2, 2.5])

doc.add_heading('Cuándo Escalar Automáticamente', level=2)
add_bullet(doc, 'Un proyecto lleva más de 48 horas bloqueado sin resolución → Nivel 2.')
add_bullet(doc, 'Un cliente ha contactado 3+ veces por el mismo problema sin resolución → Nivel 2.')
add_bullet(doc, 'Un cliente menciona cancelación, abogado, o publicación negativa → Nivel 3 inmediato.')
add_bullet(doc, 'Cualquier problema de seguridad o acceso no autorizado → Nivel 3 inmediato.')

doc.add_heading('Formato de Reporte de Escalamiento', level=2)
doc.add_paragraph('Todo escalamiento debe incluir:')
add_bullet(doc, 'Cliente: Nombre, producto, plan, fecha de inicio.')
add_bullet(doc, 'Problema: Descripción clara y concisa.')
add_bullet(doc, 'Acciones tomadas: Qué se ha intentado hasta ahora.')
add_bullet(doc, 'Impacto: ¿Afecta ingresos? ¿Afecta reputación?')
add_bullet(doc, 'Solución propuesta: Nunca escalar sin proponer una solución.')

# ========== SECTION 9 ==========
doc.add_page_break()
doc.add_heading('SECCIÓN 9: ESTRATEGIA DE PREVENCIÓN DE CHURN', level=1)
doc.add_paragraph(
    'Retener un cliente existente cuesta 5x menos que adquirir uno nuevo. El objetivo es mantener '
    'el churn mensual por debajo del 3%. Esta sección define cómo identificar y retener clientes en riesgo.'
)

doc.add_heading('Señales de Riesgo de Cancelación', level=2)
add_table(doc, ['Señal', 'Riesgo', 'Acción Inmediata'],
    [
        ['Cliente no responde mensajes en 7+ días', 'Medio', 'Llamar directamente. No enviar más mensajes de texto.'],
        ['Uso del producto disminuye significativamente', 'Alto', 'Agendar check-in proactivo para entender por qué.'],
        ['Cliente se queja 2+ veces del mismo tema', 'Alto', 'Escalar a Nivel 2. Resolver en ≤ 24 horas.'],
        ['Cliente pregunta cómo cancelar', 'Crítico', 'Activar protocolo de retención inmediatamente.'],
        ['Pago atrasado 15+ días', 'Alto', 'Contactar para entender situación. Ofrecer plan de pago.'],
        ['Cliente menciona a un competidor', 'Alto', 'Reforzar valor diferenciador. Ofrecer mejora de servicio.'],
    ],
    [2.5, 0.8, 4.0])

doc.add_heading('Protocolo de Retención (Cuando el Cliente Quiere Cancelar)', level=2)
add_bullet(doc, 'Paso 1: Llamar personalmente — nunca manejar cancelaciones por mensaje.')
add_bullet(doc, 'Paso 2: Preguntar "¿Qué no ha funcionado como esperaba?" Escuchar sin defenderse.')
add_bullet(doc, 'Paso 3: Identificar si es problema técnico (resolverlo) o de valor percibido (demostrar ROI).')
add_bullet(doc, 'Paso 4: Ofrecer una solución de retención:')

doc.add_heading('Ofertas de Retención Autorizadas', level=3)
add_table(doc, ['Oferta', 'Cuándo Usar', 'Aprobación Requerida'],
    [
        ['1 mes gratis', 'Primer intento de cancelación, cliente <3 meses de antigüedad', 'Gerente Ops'],
        ['Cambio de plan (downgrade)', 'Cliente dice que el precio es muy alto', 'Gerente Ops'],
        ['Funcionalidad adicional sin costo', 'Cliente siente que no está obteniendo suficiente valor', 'Gerente Ops + Técnico'],
        ['Descuento del 20% por 3 meses', 'Cliente valioso con alto MRR, >6 meses de antigüedad', 'CEO'],
        ['Pausa de 1 mes (sin cobro)', 'Cliente con problema temporal (vacaciones, remodelación)', 'Gerente Ops'],
    ],
    [2.0, 3.0, 1.5])

doc.add_heading('Entrevista de Salida (Si Cancela)', level=2)
doc.add_paragraph('Si el cliente cancela de todas formas, documentar:')
add_bullet(doc, '¿Cuál fue la razón principal de la cancelación?')
add_bullet(doc, '¿Qué podríamos haber hecho diferente?')
add_bullet(doc, '¿Consideraría regresar en el futuro? ¿Bajo qué condiciones?')
add_bullet(doc, '¿Nos recomendaría a otros negocios?')
p = doc.add_paragraph()
r = p.add_run('Mantener contacto trimestral con clientes cancelados — pueden regresar.')
r.italic = True; r.bold = True

# ========== SECTION 10 ==========
doc.add_page_break()
doc.add_heading('SECCIÓN 10: CHECKLIST DE CONTROL DE CALIDAD', level=1)
doc.add_paragraph(
    'Ningún producto se entrega al cliente sin pasar el control de calidad. El equipo de QA/Soporte '
    'es responsable de completar este checklist antes de invitar al cliente a la revisión.'
)

doc.add_heading('QA Universal (Aplica a Todos los Productos)', level=2)
qa_items = [
    '☐  El producto funciona correctamente en móvil (Chrome/Safari)',
    '☐  El producto funciona en desktop (Chrome/Firefox/Edge)',
    '☐  Toda la información del cliente está correcta (nombre, logo, precios, contacto)',
    '☐  No hay errores ortográficos ni texto de placeholder visible',
    '☐  Los enlaces funcionan (no hay links rotos)',
    '☐  El diseño es consistente con la marca del cliente',
    '☐  El producto se carga en menos de 3 segundos',
    '☐  Las notificaciones/emails de prueba se reciben correctamente',
]
for item in qa_items:
    add_bullet(doc, item)

doc.add_heading('QA Específico por Producto', level=2)
qa_products = [
    ('AlToque App', [
        '☐  Todos los productos tienen foto, nombre y precio',
        '☐  Pedido de prueba se completa exitosamente',
        '☐  Notificación de pedido llega al WhatsApp del cliente',
        '☐  Categorías de productos están organizadas correctamente',
    ]),
    ('Desarrollo Web', [
        '☐  SSL activo (candado verde en el navegador)',
        '☐  Formulario de contacto envía correctamente',
        '☐  Sistema de reservas funciona (si aplica)',
        '☐  SEO básico configurado (título, meta description, Google indexing solicitado)',
        '☐  Responsive: se ve bien en iPhone SE, iPhone Pro, Android, tablet, desktop',
    ]),
    ('Vacation Rentals', [
        '☐  Calendario de disponibilidad funciona',
        '☐  Reserva de prueba se completa y confirma',
        '☐  Fotos se ven en alta calidad en la galería',
        '☐  Precios dinámicos configurados (si aplica)',
        '☐  Panel del propietario accesible',
    ]),
    ('Automatización', [
        '☐  Cada flujo se ejecuta correctamente con datos de prueba',
        '☐  Triggers se disparan en el tiempo correcto',
        '☐  Manejo de errores funciona (no falla silenciosamente)',
        '☐  Integraciones de terceros conectadas y autenticadas',
    ]),
    ('Chatbot IA', [
        '☐  Responde correctamente a las FAQ del cliente',
        '☐  Escala a humano cuando no puede responder',
        '☐  Funciona en todos los canales configurados',
        '☐  Idioma y tono son consistentes con la marca del cliente',
        '☐  Funciona fuera de horario laboral',
    ]),
]
for prod_name, items in qa_products:
    doc.add_heading(prod_name, level=3)
    for item in items:
        add_bullet(doc, item)

# ========== SECTION 11 ==========
doc.add_page_break()
doc.add_heading('SECCIÓN 11: ESTÁNDARES DE DOCUMENTACIÓN Y ARCHIVOS', level=1)
doc.add_paragraph(
    'Toda la documentación interna y de clientes debe seguir estándares consistentes para '
    'facilitar la búsqueda, el traspaso entre miembros del equipo y la auditoría.'
)

doc.add_heading('Convenciones de Nombres de Archivos', level=2)
doc.add_paragraph('Formato: [CLIENTE]_[TIPO]_[VERSION].[ext]')
doc.add_paragraph('Ejemplos:')
add_bullet(doc, 'LasBrisas_Contrato_v1.pdf')
add_bullet(doc, 'CafeElTunco_Fotos_Productos.zip')
add_bullet(doc, 'RentaPlaya_Chatbot_FAQ_v2.docx')

doc.add_heading('Estructura de Carpetas por Cliente', level=2)
folder = (
    '📁 Clientes/\n'
    '   📁 [NombreCliente]/\n'
    '      📁 01_Contrato/        → Contrato firmado, formulario de intake\n'
    '      📁 02_Materiales/      → Fotos, logos, textos recibidos del cliente\n'
    '      📁 03_Implementación/  → Archivos técnicos, configs, código\n'
    '      📁 04_QA/              → Resultados de pruebas, capturas de pantalla\n'
    '      📁 05_Comunicación/    → Historial de mensajes importantes con el cliente\n'
    '      📁 06_Facturación/     → Comprobantes de pago, facturas\n'
)
t = doc.add_table(rows=1, cols=1)
t.style = 'Table Grid'
t.rows[0].cells[0].text = folder
for p in t.rows[0].cells[0].paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.name = 'Consolas'
doc.add_paragraph()

doc.add_heading('Política de Respaldo', level=2)
add_bullet(doc, 'Todos los archivos de clientes deben estar en [PLACEHOLDER: Google Drive / SharePoint / etc.]')
add_bullet(doc, 'Respaldo automático diario de la base de datos de producción.')
add_bullet(doc, 'Los archivos locales son temporales — todo debe subirse a la nube dentro de 24 horas.')
add_bullet(doc, 'Retención mínima: 2 años después de que el cliente cancele.')

doc.add_heading('Documentos Obligatorios por Cliente', level=2)
add_table(doc, ['Documento', 'Cuándo Se Crea', 'Responsable'],
    [
        ['Formulario de Intake firmado', 'Al cierre de venta', 'Rep Comercial'],
        ['Contrato de servicio firmado', 'Al cierre de venta', 'Rep Comercial'],
        ['Checklist de QA completado', 'Antes del go-live', 'QA/Soporte'],
        ['Reporte de go-live', 'Al activar el producto', 'Gerente Ops'],
        ['Encuesta de satisfacción', '30 días post go-live', 'Soporte'],
    ],
    [2.5, 2.0, 2.0])

# ========== SECTION 12 ==========
doc.add_page_break()
doc.add_heading('SECCIÓN 12: KPIs OPERACIONALES', level=1)
doc.add_paragraph(
    'Los KPIs se revisan semanalmente en la reunión de equipo y mensualmente en el reporte al CEO. '
    'Estos indicadores están alineados con el Dashboard Financiero del Fundador.'
)

doc.add_heading('KPIs Principales', level=2)
add_table(doc, ['KPI', 'Definición', 'Meta', 'Frecuencia de Medición'],
    [
        ['MRR (Monthly Recurring Revenue)', 'Suma de todas las suscripciones activas', 'Crecimiento ≥ 10% mensual', 'Mensual'],
        ['Clientes Activos', 'Total de clientes con suscripción activa pagando', 'Crecimiento neto positivo cada mes', 'Semanal'],
        ['Tasa de Churn', 'Clientes cancelados ÷ (Activos + Cancelados del mes)', '< 3% mensual', 'Mensual'],
        ['Velocidad de Onboarding', 'Días promedio desde cierre hasta go-live', '≤ 7 días (promedio)', 'Semanal'],
        ['Cumplimiento de SLA', '% de etapas del pipeline completadas dentro del SLA', '≥ 90%', 'Semanal'],
        ['Satisfacción del Cliente (NPS)', 'Net Promoter Score en encuesta post-lanzamiento', '≥ 8/10 promedio', 'Mensual'],
        ['Tiempo de Resolución de Soporte', 'Tiempo promedio para resolver tickets de soporte', '≤ 24 horas', 'Semanal'],
        ['Ingreso Residual Total', 'Suma de comisiones residuales activas de todos los reps', 'Coincide con MRR × % comisión', 'Mensual'],
    ],
    [2.0, 2.5, 1.5, 1.2])

doc.add_heading('Dashboard Mensual — Formato de Reporte', level=2)
doc.add_paragraph('El reporte mensual al CEO debe contener:')
add_table(doc, ['Sección', 'Contenido'],
    [
        ['1. Resumen Ejecutivo', '3–5 líneas: MRR actual, clientes activos, churn, logros destacados'],
        ['2. Pipeline de Onboarding', 'Número de clientes en cada etapa, cuellos de botella'],
        ['3. Producto por Producto', 'Clientes activos, MRR y churn desglosados por producto'],
        ['4. Soporte y Calidad', 'Tickets abiertos, resueltos, tiempo promedio, NPS'],
        ['5. Capacidad del Equipo', 'Utilización actual, proyección de necesidades'],
        ['6. Alertas y Riesgos', 'Clientes en riesgo, problemas no resueltos, dependencias'],
        ['7. Próximas Acciones', 'Top 3 prioridades para el mes siguiente'],
    ],
    [2.0, 5.5])

doc.add_heading('Benchmarks por Etapa de Crecimiento', level=2)
add_table(doc, ['Etapa', 'Clientes Activos', 'MRR Objetivo', 'Equipo Requerido'],
    [
        ['Inicio (actual)', '1–25', '$500–$2,000', '1 fundador + 1 técnico + 2 reps'],
        ['Crecimiento', '25–75', '$2,000–$6,000', '1 gerente ops + 2 técnicos + 4 reps'],
        ['Escala', '75–200', '$6,000–$15,000', '1 gerente ops + 4 técnicos + 1 QA + 6 reps'],
        ['Madurez', '200+', '$15,000+', 'Equipo completo con especializaciones por producto'],
    ],
    [1.2, 1.5, 1.5, 3.3])

# ========== FOOTER ==========
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = 1  # CENTER
r = p.add_run('Documento elaborado por: Todo Directo – Equipo de Operaciones  |  Versión 1.0  |  Marzo 2026')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor.from_string('808B96')
r.italic = True

# ========== SAVE FINAL ==========
OUTPUT = r'c:\Users\12132\Desktop\DigiFacil\DigiFacil\Todo_Directo_Operations_SOP.docx'
doc.save(OUTPUT)
print(f'DONE — saved to {OUTPUT}')
